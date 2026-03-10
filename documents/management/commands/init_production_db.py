"""
Инициализация production БД: создаёт пользователей и тестовые данные при старте.
Идемпотентна — не создаёт дубли при повторных запусках.
"""
import random
from io import BytesIO
from datetime import timedelta

from django.core.management.base import BaseCommand
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.utils import timezone

from accounts.models import UserProfile
from documents.models import Document, DocumentTemplate


# ─── Обязательные пользователи ────────────────────────────────────────────────
REQUIRED_USERS = [
    {
        'username': 'admin',
        'password': 'admin',
        'first_name': 'Администратор',
        'last_name': 'Системы',
        'email': 'admin@example.com',
        'is_superuser': True,
        'is_staff': True,
        'role': 'admin',
        'position': 'Системный администратор',
        'department': 'IT отдел',
    },
    {
        'username': 'александров.алексей45',
        'password': 'Hfqcel779',
        'first_name': 'Алексей',
        'last_name': 'Александров',
        'email': 'alexey45@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'clerk',
        'position': 'Делопроизводитель',
        'department': 'Канцелярия',
    },
    {
        'username': 'александров.виктор47',
        'password': 'Hfqcel779',
        'first_name': 'Виктор',
        'last_name': 'Александров',
        'email': 'viktor47@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'manager',
        'position': 'Руководитель',
        'department': 'Управление',
    },
    {
        'username': 'александров.игорь10',
        'password': 'Hfqcel779',
        'first_name': 'Игорь',
        'last_name': 'Александров',
        'email': 'igor10@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'employee',
        'position': 'Сотрудник',
        'department': 'Отдел исполнения',
    },
]

# ─── Дополнительные тестовые пользователи ─────────────────────────────────────
EXTRA_USERS = [
    ('Иванов', 'Дмитрий', 'manager', 'Руководитель отдела', 'Финансовый отдел'),
    ('Петров', 'Сергей', 'clerk', 'Старший делопроизводитель', 'Канцелярия'),
    ('Смирнова', 'Ольга', 'employee', 'Специалист', 'Бухгалтерия'),
    ('Козлов', 'Андрей', 'employee', 'Инженер', 'IT отдел'),
    ('Новикова', 'Елена', 'employee', 'Аналитик', 'Отдел продаж'),
    ('Федоров', 'Максим', 'manager', 'Начальник отдела', 'Отдел кадров'),
    ('Волкова', 'Наталья', 'clerk', 'Делопроизводитель', 'Юридический отдел'),
    ('Морозов', 'Павел', 'employee', 'Менеджер', 'Маркетинг'),
    ('Соколова', 'Мария', 'employee', 'Консультант', 'Отдел продаж'),
    ('Лебедев', 'Роман', 'employee', 'Специалист по закупкам', 'Закупки'),
    ('Захарова', 'Татьяна', 'employee', 'Бухгалтер', 'Бухгалтерия'),
    ('Егоров', 'Николай', 'manager', 'Директор департамента', 'Производство'),
    ('Кузнецова', 'Ирина', 'employee', 'HR-специалист', 'Отдел кадров'),
    ('Васильев', 'Константин', 'employee', 'Юрист', 'Юридический отдел'),
    ('Михайлова', 'Анна', 'clerk', 'Офис-менеджер', 'Административный отдел'),
]

# ─── Шаблоны ──────────────────────────────────────────────────────────────────
TEMPLATES_DATA = [
    ('Приказ о приёме на работу', 'order', 'docx',
     'Стандартный приказ о приёме сотрудника на работу'),
    ('Приказ об отпуске', 'order', 'docx',
     'Приказ о предоставлении ежегодного оплачиваемого отпуска'),
    ('Приказ о командировке', 'order', 'docx',
     'Приказ о направлении сотрудника в командировку'),
    ('Приказ о премировании', 'order', 'xlsx',
     'Приказ о выплате премии сотрудникам'),
    ('Трудовой договор', 'contract', 'docx',
     'Стандартный трудовой договор с сотрудником'),
    ('Договор поставки', 'contract', 'docx',
     'Договор на поставку товаров и услуг'),
    ('Договор подряда', 'contract', 'docx',
     'Договор подряда с подрядчиком'),
    ('Акт приёма-передачи', 'act', 'docx',
     'Акт приёма-передачи документов или имущества'),
    ('Акт выполненных работ', 'act', 'xlsx',
     'Акт о выполнении работ по договору'),
    ('Служебная записка', 'memo', 'docx',
     'Служебная записка для внутренней переписки'),
    ('Финансовый отчёт', 'report', 'xlsx',
     'Финансовый отчёт о доходах и расходах'),
    ('Квартальный отчёт', 'report', 'docx',
     'Отчёт о деятельности подразделения за квартал'),
    ('Письмо контрагенту', 'letter', 'docx',
     'Официальное письмо внешнему контрагенту'),
    ('Заявление на отпуск', 'application', 'docx',
     'Заявление сотрудника на предоставление отпуска'),
    ('Смета расходов', 'other', 'xlsx',
     'Смета расходов по проекту или мероприятию'),
]

# ─── Плейсхолдеры для каждого шаблона ─────────────────────────────────────────
TEMPLATE_PLACEHOLDERS = {
    'Приказ о приёме на работу': [
        {'name': 'ФИО', 'label': 'ФИО сотрудника', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': True},
        {'name': 'дата_приёма', 'label': 'Дата приёма на работу', 'type': 'date', 'required': True},
        {'name': 'оклад', 'label': 'Должностной оклад (руб.)', 'type': 'text', 'required': True},
        {'name': 'испытательный_срок', 'label': 'Испытательный срок (мес.)', 'type': 'text', 'required': False, 'default': '3'},
    ],
    'Приказ об отпуске': [
        {'name': 'ФИО', 'label': 'ФИО сотрудника', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': False},
        {'name': 'дата_начала', 'label': 'Дата начала отпуска', 'type': 'date', 'required': True},
        {'name': 'дата_окончания', 'label': 'Дата окончания отпуска', 'type': 'date', 'required': True},
        {'name': 'количество_дней', 'label': 'Количество календарных дней', 'type': 'number', 'required': True},
    ],
    'Приказ о командировке': [
        {'name': 'ФИО', 'label': 'ФИО сотрудника', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'место_командировки', 'label': 'Место командировки', 'type': 'text', 'required': True},
        {'name': 'цель', 'label': 'Цель командировки', 'type': 'textarea', 'required': True},
        {'name': 'дата_отъезда', 'label': 'Дата отъезда', 'type': 'date', 'required': True},
        {'name': 'дата_возврата', 'label': 'Дата возвращения', 'type': 'date', 'required': True},
        {'name': 'количество_дней', 'label': 'Количество дней командировки', 'type': 'number', 'required': True},
    ],
    'Приказ о премировании': [
        {'name': 'ФИО', 'label': 'ФИО сотрудника', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': True},
        {'name': 'основание', 'label': 'Основание для премирования', 'type': 'textarea', 'required': True},
        {'name': 'сумма', 'label': 'Сумма премии (руб.)', 'type': 'text', 'required': True},
        {'name': 'период', 'label': 'Расчётный период', 'type': 'text', 'required': True},
    ],
    'Трудовой договор': [
        {'name': 'ФИО_работника', 'label': 'ФИО работника', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': True},
        {'name': 'оклад', 'label': 'Должностной оклад (руб.)', 'type': 'text', 'required': True},
        {'name': 'дата_начала', 'label': 'Дата начала работы', 'type': 'date', 'required': True},
        {'name': 'режим_работы', 'label': 'Режим рабочего времени', 'type': 'text', 'required': False, 'default': '5/2, 9:00–18:00'},
        {'name': 'адрес_работника', 'label': 'Адрес регистрации работника', 'type': 'text', 'required': False},
    ],
    'Договор поставки': [
        {'name': 'контрагент', 'label': 'Наименование контрагента', 'type': 'text', 'required': True},
        {'name': 'ИНН_контрагента', 'label': 'ИНН контрагента', 'type': 'text', 'required': True},
        {'name': 'предмет_договора', 'label': 'Предмет договора (наименование товаров)', 'type': 'textarea', 'required': True},
        {'name': 'сумма_договора', 'label': 'Сумма договора (руб.)', 'type': 'text', 'required': True},
        {'name': 'срок_поставки', 'label': 'Срок поставки (дней)', 'type': 'text', 'required': True},
        {'name': 'условия_оплаты', 'label': 'Условия оплаты', 'type': 'text', 'required': False, 'default': '100% предоплата'},
    ],
    'Договор подряда': [
        {'name': 'подрядчик', 'label': 'Наименование подрядчика', 'type': 'text', 'required': True},
        {'name': 'ИНН_подрядчика', 'label': 'ИНН подрядчика', 'type': 'text', 'required': False},
        {'name': 'предмет_работ', 'label': 'Предмет и объём работ', 'type': 'textarea', 'required': True},
        {'name': 'стоимость', 'label': 'Стоимость работ (руб.)', 'type': 'text', 'required': True},
        {'name': 'срок_начала', 'label': 'Дата начала работ', 'type': 'date', 'required': True},
        {'name': 'срок_окончания', 'label': 'Дата окончания работ', 'type': 'date', 'required': True},
    ],
    'Акт приёма-передачи': [
        {'name': 'передаёт', 'label': 'Кто передаёт (ФИО, должность)', 'type': 'text', 'required': True},
        {'name': 'принимает', 'label': 'Кто принимает (ФИО, должность)', 'type': 'text', 'required': True},
        {'name': 'перечень', 'label': 'Перечень передаваемых документов / имущества', 'type': 'textarea', 'required': True},
        {'name': 'количество', 'label': 'Количество позиций', 'type': 'number', 'required': False},
        {'name': 'состояние', 'label': 'Состояние при передаче', 'type': 'text', 'required': False, 'default': 'Удовлетворительное'},
    ],
    'Акт выполненных работ': [
        {'name': 'исполнитель', 'label': 'Исполнитель', 'type': 'text', 'required': True},
        {'name': 'заказчик', 'label': 'Заказчик', 'type': 'text', 'required': True},
        {'name': 'номер_договора', 'label': 'Номер договора', 'type': 'text', 'required': True},
        {'name': 'наименование_работ', 'label': 'Наименование выполненных работ', 'type': 'textarea', 'required': True},
        {'name': 'стоимость', 'label': 'Стоимость работ (руб.)', 'type': 'text', 'required': True},
        {'name': 'период', 'label': 'Период выполнения работ', 'type': 'text', 'required': True},
    ],
    'Служебная записка': [
        {'name': 'от_кого', 'label': 'От кого (ФИО, должность)', 'type': 'text', 'required': True},
        {'name': 'кому', 'label': 'Кому (ФИО, должность адресата)', 'type': 'text', 'required': True},
        {'name': 'тема', 'label': 'Тема записки', 'type': 'text', 'required': True},
        {'name': 'содержание', 'label': 'Содержание записки', 'type': 'textarea', 'required': True},
    ],
    'Финансовый отчёт': [
        {'name': 'отдел', 'label': 'Наименование подразделения', 'type': 'text', 'required': True},
        {'name': 'период', 'label': 'Отчётный период', 'type': 'text', 'required': True},
        {'name': 'доходы', 'label': 'Итого доходы (руб.)', 'type': 'text', 'required': True},
        {'name': 'расходы', 'label': 'Итого расходы (руб.)', 'type': 'text', 'required': True},
        {'name': 'прибыль', 'label': 'Прибыль / убыток (руб.)', 'type': 'text', 'required': False},
        {'name': 'ответственный', 'label': 'Ответственный исполнитель', 'type': 'text', 'required': True},
    ],
    'Квартальный отчёт': [
        {'name': 'отдел', 'label': 'Наименование подразделения', 'type': 'text', 'required': True},
        {'name': 'квартал', 'label': 'Отчётный квартал (например: I кв. 2026)', 'type': 'text', 'required': True},
        {'name': 'руководитель', 'label': 'ФИО руководителя', 'type': 'text', 'required': True},
        {'name': 'ключевые_результаты', 'label': 'Ключевые результаты периода', 'type': 'textarea', 'required': True},
        {'name': 'проблемы', 'label': 'Проблемы и риски', 'type': 'textarea', 'required': False},
        {'name': 'план_следующего', 'label': 'План на следующий период', 'type': 'textarea', 'required': False},
    ],
    'Письмо контрагенту': [
        {'name': 'адресат', 'label': 'Адресат (организация)', 'type': 'text', 'required': True},
        {'name': 'контактное_лицо', 'label': 'Контактное лицо', 'type': 'text', 'required': False},
        {'name': 'тема_письма', 'label': 'Тема письма', 'type': 'text', 'required': True},
        {'name': 'содержание', 'label': 'Содержание письма', 'type': 'textarea', 'required': True},
        {'name': 'от_кого', 'label': 'Подписант (ФИО, должность)', 'type': 'text', 'required': True},
    ],
    'Заявление на отпуск': [
        {'name': 'ФИО', 'label': 'ФИО заявителя', 'type': 'text', 'required': True},
        {'name': 'должность', 'label': 'Должность', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': False},
        {'name': 'тип_отпуска', 'label': 'Вид отпуска', 'type': 'text', 'required': True, 'default': 'ежегодный оплачиваемый'},
        {'name': 'дата_начала', 'label': 'Дата начала отпуска', 'type': 'date', 'required': True},
        {'name': 'дата_окончания', 'label': 'Дата окончания отпуска', 'type': 'date', 'required': True},
        {'name': 'количество_дней', 'label': 'Количество календарных дней', 'type': 'number', 'required': True},
    ],
    'Смета расходов': [
        {'name': 'проект', 'label': 'Наименование проекта / мероприятия', 'type': 'text', 'required': True},
        {'name': 'отдел', 'label': 'Подразделение', 'type': 'text', 'required': True},
        {'name': 'период', 'label': 'Период', 'type': 'text', 'required': True},
        {'name': 'статья_1', 'label': 'Статья расходов 1', 'type': 'text', 'required': False},
        {'name': 'сумма_1', 'label': 'Сумма по статье 1 (руб.)', 'type': 'text', 'required': False},
        {'name': 'статья_2', 'label': 'Статья расходов 2', 'type': 'text', 'required': False},
        {'name': 'сумма_2', 'label': 'Сумма по статье 2 (руб.)', 'type': 'text', 'required': False},
        {'name': 'итого', 'label': 'Итоговая сумма (руб.)', 'type': 'text', 'required': True},
    ],
}


# ─── Генераторы файлов ────────────────────────────────────────────────────────

def _make_docx(title, body_text, placeholders=None):
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    placeholders = placeholders or []
    doc = DocxDoc()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    if placeholders:
        doc.add_heading('Поля для заполнения', level=2)
        for ph in placeholders:
            marker = '{{' + ph['name'] + '}}'
            label = ph.get('label', ph['name'])
            doc.add_paragraph(f'{label}: {marker}', style='List Bullet')
        doc.add_paragraph()

    doc.add_heading('Раздел 1. Основные сведения', level=2)
    for ph in placeholders:
        marker = '{{' + ph['name'] + '}}'
        label = ph.get('label', ph['name'])
        doc.add_paragraph(f'{label}: {marker}')

    if not placeholders:
        doc.add_paragraph('ФИО исполнителя: {{ФИО}}')
        doc.add_paragraph('Должность: {{должность}}')
        doc.add_paragraph('Отдел: {{отдел}}')

    doc.add_heading('Раздел 2. Содержание', level=2)
    doc.add_paragraph(body_text)
    doc.add_paragraph()
    doc.add_paragraph('Подпись: _______________')
    doc.add_paragraph('Дата подписания: _______________')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_xlsx(title, placeholders=None):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    placeholders = placeholders or []
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]

    headers = ['Поле', 'Плейсхолдер', 'Значение', 'Примечание']

    if placeholders:
        rows = [
            (ph.get('label', ph['name']), '{{' + ph['name'] + '}}', '', 'Обязательно' if ph.get('required') else 'Необязательно')
            for ph in placeholders
        ]
    else:
        rows = [
            ('ФИО', '{{ФИО}}', '', 'Заполнить'),
            ('Должность', '{{должность}}', '', 'Заполнить'),
            ('Отдел', '{{отдел}}', '', 'Заполнить'),
            ('Дата', '{{дата}}', '', 'Заполнить'),
            ('Номер', '{{номер}}', '', 'Заполнить'),
        ]

    ws.merge_cells('A1:D1')
    tc = ws['A1']
    tc.value = title
    tc.font = Font(bold=True, size=14, color='1E40AF')
    tc.alignment = Alignment(horizontal='center')
    ws.row_dimensions[1].height = 28

    fill = PatternFill(fill_type='solid', fgColor='1E40AF')
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = fill
        cell.alignment = Alignment(horizontal='center')

    for r, row in enumerate(rows, 3):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    ws.column_dimensions['A'].width = 22
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 22
    ws.column_dimensions['D'].width = 16

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def _generate_template_file(name, fmt, placeholders=None):
    placeholders = placeholders or []
    body = (
        f'Настоящий документ «{name}» составлен в соответствии с '
        'внутренними регламентами организации и является обязательным '
        'для исполнения всеми сотрудниками соответствующего подразделения.'
    )
    safe_name = name.replace(' ', '_').replace('/', '-').replace('«', '').replace('»', '')
    if fmt == 'docx':
        data = _make_docx(name, body, placeholders)
        filename = f'{safe_name}.docx'
    elif fmt == 'xlsx':
        data = _make_xlsx(name, placeholders)
        filename = f'{safe_name}.xlsx'
    else:
        data = _make_docx(name, body, placeholders)
        filename = f'{safe_name}.docx'
    return ContentFile(data), filename


# ─── Management command ───────────────────────────────────────────────────────

class Command(BaseCommand):
    help = 'Инициализация production БД: пользователи + тестовые данные'

    def handle(self, *args, **options):
        self._ensure_users()
        # Всегда восстанавливаем файлы шаблонов (нужно при каждом редеплое на Railway,
        # т.к. эфемерное хранилище сбрасывается, но PostgreSQL-записи остаются)
        self._seed_templates()
        if Document.objects.count() == 0:
            self.stdout.write('  БД пуста — загружаю тестовые данные...')
            self._seed_documents()
            self.stdout.write(self.style.SUCCESS('✅ Тестовые данные загружены'))
        else:
            self.stdout.write('  Документы уже есть — пропускаю сидинг.')

    def _ensure_users(self):
        for u in REQUIRED_USERS:
            user, created = User.objects.get_or_create(username=u['username'])
            user.set_password(u['password'])
            user.first_name = u['first_name']
            user.last_name = u['last_name']
            user.email = u['email']
            user.is_superuser = u['is_superuser']
            user.is_staff = u['is_staff']
            user.save()
            profile, _ = UserProfile.objects.get_or_create(user=user)
            profile.role = u['role']
            profile.position = u['position']
            profile.department = u['department']
            profile.save()
            self.stdout.write(
                self.style.SUCCESS(
                    f'  ✅ {"Создан" if created else "OK"}: {u["username"]} ({u["role"]})'
                )
            )

        idx = 1
        for last, first, role, position, dept in EXTRA_USERS:
            username = f'{last.lower()}.{first.lower()}{idx}'
            idx += 1
            user, created = User.objects.get_or_create(username=username)
            if created:
                user.set_password('Test1234!')
                user.first_name = first
                user.last_name = last
                user.email = f'{username}@example.com'
                user.save()
                profile, _ = UserProfile.objects.get_or_create(user=user)
                profile.role = role
                profile.position = position
                profile.department = dept
                profile.save()

    def _seed_templates(self):
        import os
        for name, type_, fmt, desc in TEMPLATES_DATA:
            tpl, created = DocumentTemplate.objects.get_or_create(name=name)
            tpl.type = type_
            tpl.file_format = fmt
            tpl.description = desc
            tpl.is_active = True
            phs = TEMPLATE_PLACEHOLDERS.get(name, [])
            tpl.placeholders = phs
            tpl.html_template = f'<h1>{name}</h1><p>{{{{ФИО}}}}, {{{{отдел}}}}</p>'

            # Файл генерируем только если он отсутствует на диске
            file_missing = (
                not tpl.template_file or
                not tpl.template_file.name or
                not os.path.exists(tpl.template_file.path)
            )
            if file_missing:
                file_content, filename = _generate_template_file(name, fmt, phs)
                tpl.template_file.save(filename, file_content, save=False)
                self.stdout.write(f'  📄 Файл пересоздан: {name} ({fmt})')
            tpl.save()
            if created:
                self.stdout.write(f'  📄 Новый шаблон: {name} ({fmt}, {len(phs)} плейсхолд.)')

    def _seed_documents(self):
        all_users = list(User.objects.all())
        templates = list(DocumentTemplate.objects.all())

        titles = [
            'Приказ о приёме на работу {name}',
            'Приказ об отпуске {name}',
            'Приказ о премировании сотрудников {dept}',
            'Приказ о командировке {name}',
            'Договор № {num} с ООО «{company}»',
            'Акт приёма-передачи № {num}',
            'Служебная записка от {dept}',
            'Докладная записка о закупке оборудования',
            'Письмо в адрес ООО «{company}»',
            'Финансовый отчёт за {period}',
            'Квартальный отчёт {dept} за {period}',
            'Заявление на отпуск {name}',
            'Трудовой договор с {name}',
            'Смета расходов {dept} на {period}',
            'Акт выполненных работ по договору № {num}',
        ]
        companies = ['Альфа', 'Бета', 'Прогресс', 'Инновация', 'Развитие', 'Омега']
        periods = ['январь 2026', 'февраль 2026', 'март 2026', 'I кв. 2026']
        statuses = ['draft', 'sent_for_approval', 'coordination', 'approved',
                    'execution', 'rejected', 'returned', 'archived']
        weights = [0.10, 0.15, 0.15, 0.25, 0.15, 0.05, 0.05, 0.10]

        for _ in range(60):
            creator = random.choice(all_users)
            tpl = random.choice(templates) if templates else None
            dept = getattr(getattr(creator, 'profile', None), 'department', 'Отдел')

            title_fmt = random.choice(titles)
            title = title_fmt.format(
                name=f'{creator.first_name} {creator.last_name}',
                dept=dept,
                num=f'{random.randint(100, 999)}/2026',
                company=random.choice(companies),
                period=random.choice(periods),
            )

            days_ago = random.randint(0, 180)
            created_at = timezone.now() - timedelta(days=days_ago)
            status = random.choices(statuses, weights=weights)[0]
            assigned = random.choice(all_users) if status != 'draft' else None

            doc = Document.objects.create(
                title=title,
                template=tpl,
                status=status,
                created_by=creator,
                assigned_to=assigned,
                content=(
                    f'Содержание документа: {title}.\n\n'
                    f'Составлен {created_at.strftime("%d.%m.%Y")}.'
                ),
                deadline=(
                    (created_at + timedelta(days=random.randint(5, 30))).date()
                    if random.random() > 0.4 else None
                ),
            )
            Document.objects.filter(pk=doc.pk).update(created_at=created_at)

        self.stdout.write('  📝 Создано 60 документов')
