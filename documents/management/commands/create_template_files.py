"""
Команда для создания реальных файлов шаблонов документов (.docx, .xlsx, .pdf)
с плейсхолдерами и полноценным содержимым, а также записей в БД.
"""
import os
from django.core.management.base import BaseCommand
from django.conf import settings


class Command(BaseCommand):
    help = 'Создаёт файлы шаблонов DOCX/XLSX/PDF и регистрирует их в базе данных'

    # ─────────────────────────────────────── helpers ────────────────────────
    def _media_path(self, *parts):
        path = os.path.join(str(settings.MEDIA_ROOT), 'templates', 'files', *parts)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        return path

    # ══════════════════════════════════════ DOCX ═════════════════════════════
    def _make_docx_prikaz(self, path):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Поля страницы
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=14, italic=False, spacing_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(spacing_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(text)
            run.bold   = bold
            run.italic = italic
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            p.alignment = align
            return p

        para('{{компания}}', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
        para('', size=6, spacing_after=0)
        para('ПРИКАЗ', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        para('№ {{номер_приказа}}  от  {{дата}}  г.', align=WD_ALIGN_PARAGRAPH.CENTER)
        para('', size=6, spacing_after=0)
        para('О приёме на работу', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
        para('', size=6, spacing_after=0)

        p_main = doc.add_paragraph()
        p_main.paragraph_format.space_after  = Pt(6)
        p_main.paragraph_format.line_spacing = Pt(18)
        p_main.paragraph_format.first_line_indent = Cm(1.25)
        run = p_main.add_run(
            'На основании трудового договора № {{номер_договора}} от {{дата_договора}} '
            'и в соответствии со статьёй 68 Трудового кодекса Российской Федерации,'
        )
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        para('ПРИКАЗЫВАЮ:', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        para('', size=6, spacing_after=0)

        items = [
            '1. Принять {{фамилия}} {{имя}} {{отчество}} на должность {{должность}} в {{отдел}} с {{дата_начала}}.',
            '2. Установить должностной оклад в размере {{оклад}} рублей в месяц согласно штатному расписанию.',
            '3. Установить испытательный срок продолжительностью {{испытательный_срок}} месяца(-ев).',
            '4. Ознакомить работника с правилами внутреннего трудового распорядка, должностной инструкцией, '
            'коллективным договором и иными локальными нормативными актами, непосредственно связанными '
            'с трудовой деятельностью работника, под роспись.',
            '5. Менеджеру по персоналу — оформить личное дело работника в установленном порядке.',
            '6. Бухгалтерии — начислять заработную плату в соответствии с настоящим приказом.',
            '7. Контроль за исполнением настоящего приказа возложить на {{руководитель_отдела}}.',
        ]
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.left_indent  = Cm(0)
            run = p.add_run(item)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'

        para('', size=6, spacing_after=0)

        # Подписи
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].clear()
                for border_name in ('top', 'bottom', 'left', 'right'):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                    tcPr.append(tcBorders)

        def cell_text(cell, text, bold=False):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
            run.bold = bold

        cell_text(table.cell(0, 0), 'Руководитель организации', bold=True)
        cell_text(table.cell(0, 1), '{{руководитель}}')
        cell_text(table.cell(1, 0), 'С приказом ознакомлен(а)', bold=True)
        cell_text(table.cell(1, 1), '{{сотрудник_подпись}}')

        doc.save(path)
        return path

    def _make_docx_dogovor(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def h(text, size=14, bold=False, center=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            run.bold = bold
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def body(text):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            return p

        h('ТРУДОВОЙ ДОГОВОР', size=16, bold=True, center=True)
        h('№ {{номер_договора}}', center=True)
        h('г. {{город}}', center=True)
        h('"___" ____________ {{год}} г.', center=True)
        h('')

        body(
            '{{компания}}, именуемое в дальнейшем «Работодатель», в лице {{должность_руководителя}} '
            '{{руководитель}}, действующего на основании {{основание}}, с одной стороны, и '
            '{{фамилия}} {{имя}} {{отчество}}, именуемый(ая) в дальнейшем «Работник», с другой '
            'стороны, заключили настоящий трудовой договор о нижеследующем:'
        )
        h('')

        sections_text = [
            ('1. ПРЕДМЕТ ДОГОВОРА', [
                '1.1. Работодатель обязуется предоставить Работнику работу по обусловленной '
                'трудовой функции — {{должность}} в {{отдел}}, обеспечить условия труда, '
                'предусмотренные трудовым законодательством, своевременно и в полном объёме '
                'выплачивать Работнику заработную плату.',
                '1.2. Работник обязуется лично выполнять работу по должности {{должность}}, '
                'соблюдать правила внутреннего трудового распорядка, действующие у Работодателя.',
                '1.3. Место работы: {{адрес_работы}}.',
                '1.4. Настоящий договор является договором по {{тип_занятости}} (основному '
                'месту работы / по совместительству).',
            ]),
            ('2. СРОК ДОГОВОРА', [
                '2.1. Настоящий трудовой договор заключён на {{срок_договора}}.',
                '2.2. Дата начала работы: {{дата_начала}}.',
                '2.3. Работнику устанавливается испытательный срок продолжительностью '
                '{{испытательный_срок}} месяца(-ев). В период испытания на Работника '
                'распространяются положения трудового законодательства.',
            ]),
            ('3. УСЛОВИЯ ОПЛАТЫ ТРУДА', [
                '3.1. За выполнение трудовых обязанностей Работнику устанавливается '
                'должностной оклад в размере {{оклад}} ({{оклад_прописью}}) рублей в месяц.',
                '3.2. Заработная плата выплачивается два раза в месяц: {{дата_выплаты_1}} и '
                '{{дата_выплаты_2}} числа каждого месяца.',
                '3.3. Работодатель вправе устанавливать надбавки, доплаты, премии и иные '
                'поощрительные выплаты в порядке, установленном локальными нормативными актами.',
            ]),
            ('4. РЕЖИМ РАБОЧЕГО ВРЕМЕНИ', [
                '4.1. Работнику устанавливается {{режим_работы}} рабочая неделя '
                'продолжительностью {{часы_в_неделю}} часов.',
                '4.2. Время начала работы: {{время_начала}}. Время окончания: {{время_окончания}}.',
                '4.3. Перерыв для отдыха и питания: {{обеденный_перерыв}} минут.',
                '4.4. Работнику предоставляется ежегодный основной оплачиваемый отпуск '
                'продолжительностью {{отпуск}} календарных дней.',
            ]),
            ('5. ПРАВА И ОБЯЗАННОСТИ СТОРОН', [
                '5.1. Работник обязан: добросовестно исполнять трудовые обязанности; соблюдать '
                'правила внутреннего трудового распорядка; соблюдать трудовую дисциплину; '
                'бережно относиться к имуществу Работодателя; незамедлительно сообщать '
                'Работодателю о любой ситуации, угрожающей жизни и здоровью людей.',
                '5.2. Работодатель обязан: предоставлять Работнику работу, обусловленную '
                'трудовым договором; обеспечивать безопасность и условия труда, соответствующие '
                'государственным нормативным требованиям охраны труда; выплачивать в полном '
                'размере причитающуюся Работнику заработную плату.',
            ]),
            ('6. ОТВЕТСТВЕННОСТЬ СТОРОН', [
                '6.1. Стороны несут ответственность за неисполнение или ненадлежащее исполнение '
                'своих обязательств по настоящему договору в соответствии с законодательством '
                'Российской Федерации.',
                '6.2. Материальная ответственность Работника наступает в порядке, предусмотренном '
                'главой 39 Трудового кодекса Российской Федерации.',
            ]),
            ('7. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ', [
                '7.1. Условия настоящего трудового договора могут быть изменены только по '
                'соглашению сторон и в письменной форме.',
                '7.2. Настоящий договор составлен в двух экземплярах, имеющих одинаковую '
                'юридическую силу, один из которых хранится у Работодателя, другой — у Работника.',
                '7.3. Во всём, что не предусмотрено настоящим договором, стороны руководствуются '
                'действующим законодательством Российской Федерации.',
            ]),
        ]

        for title, paragraphs in sections_text:
            h(title, bold=True)
            for para_text in paragraphs:
                body(para_text)
            h('')

        h('8. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', bold=True)
        h('')

        table = doc.add_table(rows=6, cols=2)
        data = [
            ('РАБОТОДАТЕЛЬ', 'РАБОТНИК'),
            ('{{компания}}', '{{фамилия}} {{имя}} {{отчество}}'),
            ('ИНН: {{инн_компании}}', 'Паспорт: серия {{серия_паспорта}} № {{номер_паспорта}}'),
            ('Адрес: {{адрес_компании}}', 'Адрес: {{адрес_работника}}'),
            ('Подпись: ________________', 'Подпись: ________________'),
            ('{{руководитель}}', 'Экземпляр договора получен: "___" ________ {{год}} г.'),
        ]
        for r_idx, (left, right) in enumerate(data):
            for c_idx, text in enumerate([left, right]):
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if r_idx == 0:
                    run.bold = True

        doc.save(path)
        return path

    def _make_docx_akt(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def h(text, bold=False, center=False, size=14):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            run.bold = bold
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def body(text):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            return p

        h('АКТ ВЫПОЛНЕННЫХ РАБОТ (ОКАЗАННЫХ УСЛУГ)', bold=True, center=True, size=15)
        h('№ {{номер_акта}} от {{дата}} г.', center=True)
        h('')

        body(
            '{{компания_заказчик}}, именуемое далее «Заказчик», в лице {{должность_заказчика}} '
            '{{руководитель_заказчика}}, действующего на основании {{основание_заказчика}}, '
            'с одной стороны, и {{компания_исполнитель}}, именуемое далее «Исполнитель», в лице '
            '{{должность_исполнителя}} {{руководитель_исполнителя}}, действующего на основании '
            '{{основание_исполнителя}}, с другой стороны, составили настоящий акт о следующем:'
        )
        h('')
        h('1. Исполнитель выполнил, а Заказчик принял следующие работы (услуги):', bold=True)
        h('')

        # Таблица работ
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(['№', 'Наименование работ/услуг', 'Ед. изм.', 'Кол-во', 'Сумма, руб.']):
            p = hdr[i].paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i in range(1, 4):
            row = tbl.add_row().cells
            for j, val in enumerate([str(i), '{{наименование_работы_' + str(i) + '}}',
                                      '{{единица_' + str(i) + '}}',
                                      '{{количество_' + str(i) + '}}',
                                      '{{сумма_' + str(i) + '}}']):
                p = row[j].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'

        # Итого
        row_total = tbl.add_row().cells
        row_total[3].paragraphs[0].add_run('Итого:').bold = True
        run_sum = row_total[4].paragraphs[0].add_run('{{итого}}')
        run_sum.bold = True
        run_sum.font.size = Pt(12)
        run_sum.font.name = 'Times New Roman'

        h('')
        body(
            '2. Общая стоимость выполненных работ составляет {{итого}} ({{итого_прописью}}) '
            'рублей, в том числе НДС {{ндс}} %: {{сумма_ндс}} рублей.'
        )
        body(
            '3. Вышеперечисленные работы выполнены в полном объёме, в установленные сроки. '
            'Заказчик претензий к качеству и срокам выполнения работ не имеет.'
        )
        body(
            '4. Настоящий акт составлен в двух экземплярах, имеющих одинаковую юридическую '
            'силу, по одному для каждой из сторон.'
        )
        h('')
        h('ПОДПИСИ СТОРОН:', bold=True)
        h('')

        tbl2 = doc.add_table(rows=4, cols=2)
        rows_data = [
            ('ЗАКАЗЧИК', 'ИСПОЛНИТЕЛЬ'),
            ('{{компания_заказчик}}', '{{компания_исполнитель}}'),
            ('{{должность_заказчика}} {{руководитель_заказчика}}',
             '{{должность_исполнителя}} {{руководитель_исполнителя}}'),
            ('Подпись: _______ / М.П.', 'Подпись: _______ / М.П.'),
        ]
        for r_i, (l, r) in enumerate(rows_data):
            for c_i, txt in enumerate([l, r]):
                cell = tbl2.cell(r_i, c_i)
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if r_i == 0:
                    run.bold = True

        doc.save(path)
        return path

    # ══════════════════════════════════════ XLSX ═════════════════════════════
    def _make_xlsx_schet(self, path):
        from openpyxl import Workbook
        from openpyxl.styles import (Font, Alignment, PatternFill, Border, Side,
                                     numbers as num_fmt)
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = 'Счёт'

        thin = Side(style='thin')
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def cell(row, col, value, bold=False, center=False, fill=None, font_size=11, wrap=False, number_format=None):
            c = ws.cell(row=row, column=col, value=value)
            c.font = Font(name='Times New Roman', size=font_size, bold=bold)
            c.alignment = Alignment(
                horizontal='center' if center else 'left',
                vertical='center',
                wrap_text=wrap,
            )
            c.border = border
            if fill:
                c.fill = PatternFill('solid', fgColor=fill)
            if number_format:
                c.number_format = number_format
            return c

        # Устанавливаем ширины колонок
        col_widths = {1: 6, 2: 45, 3: 10, 4: 14, 5: 16}
        for col, w in col_widths.items():
            ws.column_dimensions[get_column_letter(col)].width = w

        # Шапка
        ws.row_dimensions[1].height = 25
        ws.merge_cells('A1:E1')
        c = ws['A1']
        c.value = 'СЧЁТ НА ОПЛАТУ № {{номер_счета}} от {{дата}} г.'
        c.font = Font(name='Times New Roman', size=14, bold=True)
        c.alignment = Alignment(horizontal='center', vertical='center')
        c.fill = PatternFill('solid', fgColor='2563EB')
        c.font = Font(name='Times New Roman', size=14, bold=True, color='FFFFFF')

        r = 2
        ws.row_dimensions[r].height = 18
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'Поставщик:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{компания_поставщик}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 3
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'ИНН/КПП:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{инн_поставщик}} / {{кпп_поставщик}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 4
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'Банк поставщика:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{банк_поставщик}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 5
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'р/с:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{расчетный_счет}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 6
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'БИК:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{бик}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 7
        ws.row_dimensions[r].height = 5

        r = 8
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'Покупатель:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{компания_покупатель}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 9
        ws.merge_cells(f'A{r}:B{r}')
        ws[f'A{r}'].value = 'Адрес покупателя:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells(f'C{r}:E{r}')
        ws[f'C{r}'].value = '{{адрес_покупателя}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)

        r = 10
        ws.row_dimensions[r].height = 5

        # Заголовок таблицы
        r = 11
        ws.row_dimensions[r].height = 30
        for col, (txt, center) in enumerate([
            ('№', True), ('Наименование товара/услуги', False),
            ('Ед.изм.', True), ('Кол-во', True), ('Сумма, руб.', True)
        ], 1):
            c = ws.cell(row=r, column=col, value=txt)
            c.font = Font(name='Times New Roman', size=11, bold=True)
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.border = border
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Строки товаров
        for i in range(1, 6):
            r = 11 + i
            ws.row_dimensions[r].height = 20
            for col, val in enumerate([
                str(i),
                f'{{{{наименование_{i}}}}}',
                f'{{{{единица_{i}}}}}',
                f'{{{{количество_{i}}}}}',
                f'{{{{сумма_{i}}}}}',
            ], 1):
                c = ws.cell(row=r, column=col, value=val)
                c.font = Font(name='Times New Roman', size=11)
                c.border = border
                c.alignment = Alignment(
                    horizontal='center' if col in (1, 3, 4, 5) else 'left',
                    vertical='center'
                )

        # Итоги
        r = 17
        ws.merge_cells(f'A{r}:D{r}')
        ws[f'A{r}'].value = 'Итого без НДС:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'A{r}'].alignment = Alignment(horizontal='right', vertical='center')
        ws[f'E{r}'].value = '{{сумма_без_ндс}}'
        ws[f'E{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'E{r}'].border = border
        ws[f'E{r}'].alignment = Alignment(horizontal='center', vertical='center')

        r = 18
        ws.merge_cells(f'A{r}:D{r}')
        ws[f'A{r}'].value = 'НДС {{ставка_ндс}}%:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'A{r}'].alignment = Alignment(horizontal='right', vertical='center')
        ws[f'E{r}'].value = '{{сумма_ндс}}'
        ws[f'E{r}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r}'].border = border
        ws[f'E{r}'].alignment = Alignment(horizontal='center', vertical='center')

        r = 19
        ws.merge_cells(f'A{r}:D{r}')
        ws[f'A{r}'].value = 'ИТОГО К ОПЛАТЕ:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=12, bold=True)
        ws[f'A{r}'].fill = PatternFill('solid', fgColor='DBEAFE')
        ws[f'A{r}'].alignment = Alignment(horizontal='right', vertical='center')
        ws[f'E{r}'].value = '{{итого}}'
        ws[f'E{r}'].font = Font(name='Times New Roman', size=12, bold=True)
        ws[f'E{r}'].fill = PatternFill('solid', fgColor='DBEAFE')
        ws[f'E{r}'].border = border
        ws[f'E{r}'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[r].height = 22

        r = 20
        ws.merge_cells(f'A{r}:E{r}')
        ws[f'A{r}'].value = 'Итого к оплате (прописью): {{итого_прописью}} рублей {{копейки}} копеек'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, italic=True)
        ws[f'A{r}'].alignment = Alignment(horizontal='left', vertical='center')

        r = 21
        ws.row_dimensions[r].height = 5

        r = 22
        ws.merge_cells(f'A{r}:E{r}')
        ws[f'A{r}'].value = (
            'Оплата данного счёта означает согласие с условиями поставки товара/оказания услуг. '
            'Уведомление об оплате обязательно. '
            'Срок оплаты: {{срок_оплаты}}.'
        )
        ws[f'A{r}'].font = Font(name='Times New Roman', size=10, italic=True)
        ws[f'A{r}'].alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        ws.row_dimensions[r].height = 32

        r = 23
        ws[f'A{r}'].value = 'Руководитель:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r}'].value = '{{руководитель}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r}'].value = '____________________'
        ws[f'E{r}'].font = Font(name='Times New Roman', size=11)

        r = 24
        ws[f'A{r}'].value = 'Главный бухгалтер:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r}'].value = '{{главный_бухгалтер}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r}'].value = '____________________'
        ws[f'E{r}'].font = Font(name='Times New Roman', size=11)

        wb.save(path)
        return path

    def _make_xlsx_vedomost(self, path):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = 'Ведомость'

        thin = Side(style='thin')
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        col_widths = {1: 5, 2: 30, 3: 15, 4: 12, 5: 12, 6: 12, 7: 14, 8: 16}
        for col, w in col_widths.items():
            ws.column_dimensions[get_column_letter(col)].width = w

        ws.merge_cells('A1:H1')
        ws['A1'].value = '{{компания}}'
        ws['A1'].font = Font(name='Times New Roman', size=13, bold=True)
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 22

        ws.merge_cells('A2:H2')
        ws['A2'].value = 'РАСЧЁТНАЯ ВЕДОМОСТЬ ЗА {{месяц}} {{год}} Г.'
        ws['A2'].font = Font(name='Times New Roman', size=14, bold=True)
        ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
        ws['A2'].fill = PatternFill('solid', fgColor='2563EB')
        ws['A2'].font = Font(name='Times New Roman', size=14, bold=True, color='FFFFFF')
        ws.row_dimensions[2].height = 26

        ws.merge_cells('A3:H3')
        ws['A3'].value = 'Подразделение: {{отдел}}'
        ws['A3'].font = Font(name='Times New Roman', size=11)
        ws['A3'].alignment = Alignment(horizontal='left', vertical='center')

        # Заголовки таблицы
        headers = ['№', 'ФИО сотрудника', 'Должность', 'Оклад', 'Отработано\n(дней)',
                   'Начислено', 'Удержано\n(НДФЛ)', 'К выплате']
        r = 5
        ws.row_dimensions[r].height = 36
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=r, column=col, value=h)
            c.font = Font(name='Times New Roman', size=11, bold=True)
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.border = border
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Строки сотрудников
        for i in range(1, 8):
            row_r = 5 + i
            ws.row_dimensions[row_r].height = 20
            vals = [
                str(i),
                f'{{{{фио_{i}}}}}',
                f'{{{{должность_{i}}}}}',
                f'{{{{оклад_{i}}}}}',
                f'{{{{дней_{i}}}}}',
                f'{{{{начислено_{i}}}}}',
                f'{{{{ндфл_{i}}}}}',
                f'{{{{к_выплате_{i}}}}}',
            ]
            for col, val in enumerate(vals, 1):
                c = ws.cell(row=row_r, column=col, value=val)
                c.font = Font(name='Times New Roman', size=11)
                c.border = border
                c.alignment = Alignment(
                    horizontal='center' if col != 2 else 'left',
                    vertical='center'
                )

        # Итоговая строка
        r_total = 13
        ws.row_dimensions[r_total].height = 22
        ws.merge_cells(f'A{r_total}:C{r_total}')
        ws[f'A{r_total}'].value = 'ИТОГО:'
        ws[f'A{r_total}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'A{r_total}'].alignment = Alignment(horizontal='right', vertical='center')
        ws[f'A{r_total}'].fill = PatternFill('solid', fgColor='DBEAFE')
        for col, val in [(4, '{{итого_оклад}}'), (6, '{{итого_начислено}}'),
                         (7, '{{итого_ндфл}}'), (8, '{{итого_к_выплате}}')]:
            c = ws.cell(row=r_total, column=col, value=val)
            c.font = Font(name='Times New Roman', size=11, bold=True)
            c.border = border
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.alignment = Alignment(horizontal='center', vertical='center')

        r = 15
        ws[f'A{r}'].value = 'Руководитель:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r}'].value = '{{руководитель}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)
        ws[f'G{r}'].value = '_____________ / М.П.'
        ws[f'G{r}'].font = Font(name='Times New Roman', size=11)

        r = 16
        ws[f'A{r}'].value = 'Главный бухгалтер:'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r}'].value = '{{главный_бухгалтер}}'
        ws[f'C{r}'].font = Font(name='Times New Roman', size=11)
        ws[f'G{r}'].value = '_____________ / М.П.'
        ws[f'G{r}'].font = Font(name='Times New Roman', size=11)

        r = 17
        ws[f'A{r}'].value = f'Дата составления: {{{{дата}}}}'
        ws[f'A{r}'].font = Font(name='Times New Roman', size=11)

        wb.save(path)
        return path

    # ══════════════════════════════════════ PDF ══════════════════════════════
    def _make_pdf_zayavlenie(self, path):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import sys

        # Регистрация шрифта с поддержкой кириллицы
        font_paths = [
            'C:/Windows/Fonts/times.ttf',
            'C:/Windows/Fonts/Times New Roman.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
            '/usr/share/fonts/truetype/freefont/FreeSerif.ttf',
        ]
        font_name = 'Helvetica'  # fallback
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('TimesRu', fp))
                    font_name = 'TimesRu'
                    break
                except Exception:
                    pass

        c = Canvas(path, pagesize=A4)
        w, h = A4

        def text(x, y, s, size=12, bold=False):
            c.setFont(font_name, size)
            c.drawString(x, y, s)

        y = h - 2 * cm
        text(w / 2 - 5 * cm, y, 'ЗАЯВЛЕНИЕ', size=16)
        y -= 0.8 * cm
        text(w / 2 - 5 * cm, y, 'на отпуск', size=13)
        y -= 1.0 * cm
        text(2 * cm, y, 'Кому: {{должность_руководителя}} {{руководитель}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, '{{компания}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'От: {{фамилия}} {{имя}} {{отчество}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'Должность: {{должность}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'Подразделение: {{отдел}}', size=12)
        y -= 1.0 * cm

        lines = [
            'Прошу предоставить мне ежегодный оплачиваемый отпуск',
            'продолжительностью {{продолжительность_отпуска}} календарных дней',
            'с {{дата_начала}} по {{дата_окончания}} включительно.',
            '',
            'Основание: график отпусков на {{год}} год / личное заявление.',
            '',
            'Дата: {{дата}}',
            '',
            'Подпись: _______________________  {{фамилия}} {{имя}} {{отчество}}',
            '',
            '',
            'Отметка руководителя:',
            '',
            'Согласовано / Не согласовано (нужное подчеркнуть)',
            '',
            'Подпись руководителя: _______________________  {{руководитель}}',
            'Дата: {{дата_согласования}}',
        ]
        for line in lines:
            text(2 * cm, y, line, size=12)
            y -= 0.6 * cm

        c.save()
        return path

    def _make_pdf_memo(self, path):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_paths = [
            'C:/Windows/Fonts/times.ttf',
            'C:/Windows/Fonts/Times New Roman.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
        ]
        font_name = 'Helvetica'
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('TimesRu2', fp))
                    font_name = 'TimesRu2'
                    break
                except Exception:
                    pass

        c = Canvas(path, pagesize=A4)
        w, h = A4

        def text(x, y, s, size=12):
            c.setFont(font_name, size)
            c.drawString(x, y, s)

        y = h - 2 * cm
        text(w / 2 - 6 * cm, y, 'СЛУЖЕБНАЯ ЗАПИСКА', size=15)
        y -= 0.7 * cm
        text(w / 2 - 3.5 * cm, y, '№ {{номер}} от {{дата}} г.', size=12)
        y -= 1.0 * cm
        text(2 * cm, y, 'Кому: {{кому_должность}} {{кому_фио}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'От: {{от_должность}} {{от_фио}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'Тема: {{тема}}', size=12)
        y -= 1.0 * cm

        paras = [
            'Настоящим довожу до Вашего сведения следующее:',
            '',
            '{{текст_сообщения}}',
            '',
            'В связи с вышеизложенным прошу {{просьба}}.',
            '',
            'Приложения: {{приложения}}.',
            '',
            '',
            'С уважением,',
            '{{от_должность}}  {{от_фио}}',
            '',
            'Дата: {{дата}}    Подпись: _________________________',
        ]
        for line in paras:
            text(2 * cm, y, line, size=12)
            y -= 0.6 * cm

        c.save()
        return path

    def _make_docx_doverennost(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def h(text, bold=False, center=False, size=14, sp_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(sp_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(18)
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            run.bold = bold
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def body(text, indent=True):
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            return p

        h('ДОВЕРЕННОСТЬ', bold=True, center=True, size=16)
        h('№ {{номер_доверенности}}', center=True)
        h('')
        h('г. {{город}},  {{дата}} г.', center=True)
        h('')

        body(
            '{{компания}}, именуемое в дальнейшем «Доверитель», в лице {{должность_руководителя}} '
            '{{руководитель}}, действующего на основании {{основание}}, настоящей доверенностью уполномочивает'
        )
        body(
            '{{фамилия}} {{имя}} {{отчество}}, '
            'паспорт серии {{серия_паспорта}} № {{номер_паспорта}}, '
            'выдан {{кем_выдан}} {{дата_выдачи}}, '
            'проживающего(ую) по адресу: {{адрес}},'
        )
        body(
            'именуемого(ую) в дальнейшем «Поверенный», представлять интересы '
            '{{компания}} в {{место_представления}} по вопросам, '
            'связанным с {{предмет_доверенности}}.'
        )
        h('')
        body(
            'Поверенный вправе совершать все необходимые действия, связанные с выполнением '
            'данного поручения, в том числе: {{перечень_действий}}.'
        )
        h('')
        body(
            'Настоящая доверенность выдана сроком до {{срок_действия}} '
            'без права передоверия.'
        )
        h('')
        h('Подпись поверенного: _____________________  {{фамилия}} {{имя}} {{отчество}}')
        h('')

        tbl = doc.add_table(rows=3, cols=2)
        for r_i, (l, r) in enumerate([
            ('ДОВЕРИТЕЛЬ', ''),
            ('{{компания}}', ''),
            ('{{руководитель}}  _________________ / М.П.', ''),
        ]):
            for c_i, txt in enumerate([l, r]):
                cell = tbl.cell(r_i, c_i)
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if r_i == 0:
                    run.bold = True

        doc.save(path)
        return path

    def _make_docx_kuplia_prodazha(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def h(text, bold=False, center=False, size=14, sp=4):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(sp)
            p.paragraph_format.line_spacing = Pt(18)
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.name = 'Times New Roman'
            r.bold = bold
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def body(text):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(20)
            r = p.add_run(text)
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            return p

        h('ДОГОВОР КУПЛИ-ПРОДАЖИ', bold=True, center=True, size=16)
        h('№ {{номер_договора}}', center=True)
        h('г. {{город}},  {{дата}} г.', center=True)
        h('')

        body(
            '{{продавец}}, именуемый(ая) в дальнейшем «Продавец», и '
            '{{покупатель}}, именуемый(ая) в дальнейшем «Покупатель», '
            'заключили настоящий договор о нижеследующем:'
        )
        h('')

        sections_text = [
            ('1. ПРЕДМЕТ ДОГОВОРА', [
                '1.1. Продавец обязуется передать в собственность Покупателя, '
                'а Покупатель обязуется принять и оплатить следующий товар: {{наименование_товара}}.',
                '1.2. Количество товара: {{количество}} {{единица_измерения}}.',
                '1.3. Качество и характеристики товара: {{характеристики}}.',
                '1.4. Товар передаётся по адресу: {{адрес_передачи}}.',
            ]),
            ('2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ', [
                '2.1. Цена единицы товара: {{цена_единицы}} рублей.',
                '2.2. Общая стоимость товара составляет {{общая_стоимость}} ({{стоимость_прописью}}) рублей, '
                'в том числе НДС {{ндс}}%.',
                '2.3. Покупатель производит оплату в течение {{срок_оплаты}} банковских дней с момента '
                'подписания настоящего договора путём перечисления денежных средств на расчётный счёт Продавца.',
            ]),
            ('3. ПЕРЕДАЧА ТОВАРА', [
                '3.1. Продавец обязуется передать товар Покупателю в срок до {{дата_передачи}}.',
                '3.2. Передача товара оформляется подписанием акта приёма-передачи.',
                '3.3. Право собственности на товар переходит к Покупателю с момента подписания '
                'акта приёма-передачи.',
            ]),
            ('4. ОТВЕТСТВЕННОСТЬ СТОРОН', [
                '4.1. За нарушение сроков оплаты Покупатель уплачивает Продавцу пеню в размере '
                '{{пеня}}% от неоплаченной суммы за каждый день просрочки.',
                '4.2. За нарушение сроков поставки Продавец уплачивает Покупателю пеню в размере '
                '{{пеня}}% от стоимости непереданного товара за каждый день просрочки.',
            ]),
            ('5. ПРОЧИЕ УСЛОВИЯ', [
                '5.1. Настоящий договор вступает в силу с момента его подписания обеими сторонами.',
                '5.2. Договор составлен в двух экземплярах, имеющих равную юридическую силу.',
                '5.3. Все споры разрешаются путём переговоров, а при недостижении соглашения — '
                'в Арбитражном суде по месту нахождения Продавца.',
            ]),
        ]
        for title, paragraphs in sections_text:
            h(title, bold=True)
            for pt in paragraphs:
                body(pt)
            h('')

        h('6. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', bold=True)
        h('')
        tbl = doc.add_table(rows=5, cols=2)
        data = [
            ('ПРОДАВЕЦ', 'ПОКУПАТЕЛЬ'),
            ('{{продавец}}', '{{покупатель}}'),
            ('ИНН: {{инн_продавца}}', 'ИНН: {{инн_покупателя}}'),
            ('р/с: {{счет_продавца}}', 'р/с: {{счет_покупателя}}'),
            ('Подпись: _____________ / М.П.', 'Подпись: _____________ / М.П.'),
        ]
        for r_i, (l, r) in enumerate(data):
            for c_i, txt in enumerate([l, r]):
                cell = tbl.cell(r_i, c_i)
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if r_i == 0:
                    run.bold = True

        doc.save(path)
        return path

    def _make_docx_uvolnenie(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def para(text, bold=False, center=False, size=14, sp_after=6, indent=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(sp_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = Pt(18)
            if indent:
                p.paragraph_format.first_line_indent = Cm(1.25)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        para('{{компания}}', bold=True, center=True)
        para('')
        para('ПРИКАЗ', bold=True, center=True, size=16)
        para('№ {{номер_приказа}}  от  {{дата}}  г.', center=True)
        para('')
        para('О прекращении (расторжении) трудового договора с работником (увольнении)',
             bold=True, center=True)
        para('')

        para(
            'На основании {{основание_увольнения}}, статья {{статья_тк}} Трудового кодекса '
            'Российской Федерации,',
            indent=True
        )
        para('ПРИКАЗЫВАЮ:', bold=True)
        para('')

        items = [
            '1. Прекратить действие трудового договора № {{номер_договора}} от {{дата_договора}}, '
            'уволить {{дата_увольнения}}:',
            '   Фамилия, имя, отчество: {{фамилия}} {{имя}} {{отчество}}',
            '   Табельный номер: {{табельный_номер}}',
            '   Структурное подразделение: {{отдел}}',
            '   Должность (специальность, профессия): {{должность}}',
            '   Основание: {{основание_увольнения}}',
            '',
            '2. Бухгалтерии произвести расчёт с работником.',
            '3. Отделу кадров оформить документы на увольнение в установленном порядке.',
            '4. Ответственному лицу принять материальные ценности и оборудование.',
        ]
        for item in items:
            para(item)

        para('')
        para(
            'Основание: {{документ_основание}} от {{дата_документа}}.',
            indent=True
        )
        para('')

        tbl = doc.add_table(rows=4, cols=2)
        sign_data = [
            ('Директор:', '{{руководитель}}'),
            ('Подпись:', '______________________'),
            ('С приказом ознакомлен(а):', '{{фамилия}} {{имя}} {{отчество}}'),
            ('Дата:', '{{дата_ознакомления}}'),
        ]
        for r_i, (l, r) in enumerate(sign_data):
            for c_i, txt in enumerate([l, r]):
                cell = tbl.cell(r_i, c_i)
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                if c_i == 0:
                    run.bold = True

        doc.save(path)
        return path

    def _make_docx_protokol(self, path):
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(1.5)

        def h(text, bold=False, center=False, size=14, sp=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(sp)
            p.paragraph_format.line_spacing = Pt(18)
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.name = 'Times New Roman'
            r.bold = bold
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def body(text, indent=True):
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(20)
            r = p.add_run(text)
            r.font.size = Pt(14)
            r.font.name = 'Times New Roman'
            return p

        h('{{компания}}', bold=True, center=True)
        h('')
        h('ПРОТОКОЛ СОВЕЩАНИЯ', bold=True, center=True, size=16)
        h('№ {{номер_протокола}}', center=True)
        h('Дата: {{дата}}', center=True)
        h('Место проведения: {{место}}', center=True)
        h('Начало: {{время_начала}}   Окончание: {{время_окончания}}', center=True)
        h('')
        h('Председатель: {{председатель}}', bold=False)
        h('Секретарь: {{секретарь}}', bold=False)
        h('')
        h('Присутствовали:', bold=True)
        for i in range(1, 6):
            h(f'{i}. {{{{участник_{i}}}}}  — {{{{должность_участника_{i}}}}}')
        h('')
        h('ПОВЕСТКА ДНЯ:', bold=True)
        for i in range(1, 4):
            h(f'{i}. {{{{вопрос_{i}}}}}')
        h('')
        for i in range(1, 4):
            h(f'{i}. СЛУШАЛИ:', bold=True)
            body(f'{{{{докладчик_{i}}}}} — {{{{содержание_доклада_{i}}}}}')
            h(f'РЕШИЛИ:', bold=True)
            body(f'{{{{решение_{i}}}}}')
            h(f'Ответственный: {{{{ответственный_{i}}}}}   Срок: {{{{срок_{i}}}}}')
            h('')

        h('Председатель:  _____________________  {{председатель}}')
        h('Секретарь:     _____________________  {{секретарь}}')

        doc.save(path)
        return path

    # ─── XLSX ────────────────────────────────────────────────────────────────
    def _make_xlsx_tabel(self, path):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = 'Табель'

        thin = Side(style='thin')
        brd = Border(left=thin, right=thin, top=thin, bottom=thin)

        # Ширины
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 28
        ws.column_dimensions['C'].width = 15
        for i in range(4, 36):           # дни 1-31 + итого
            ws.column_dimensions[get_column_letter(i)].width = 4
        ws.column_dimensions[get_column_letter(36)].width = 8
        ws.column_dimensions[get_column_letter(37)].width = 8

        # Заголовок
        ws.merge_cells('A1:AK1')
        ws['A1'].value = '{{компания}}'
        ws['A1'].font = Font(name='Times New Roman', size=13, bold=True)
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 20

        ws.merge_cells('A2:AK2')
        ws['A2'].value = 'ТАБЕЛЬ УЧЁТА РАБОЧЕГО ВРЕМЕНИ ЗА {{месяц}} {{год}} Г.'
        ws['A2'].font = Font(name='Times New Roman', size=13, bold=True, color='FFFFFF')
        ws['A2'].fill = PatternFill('solid', fgColor='2563EB')
        ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[2].height = 24

        ws['A3'].value = 'Подразделение:'
        ws['A3'].font = Font(name='Times New Roman', size=11, bold=True)
        ws.merge_cells('B3:F3')
        ws['B3'].value = '{{отдел}}'
        ws['B3'].font = Font(name='Times New Roman', size=11)

        # Заголовки таблицы (строка 5)
        headers_row = 5
        ws.row_dimensions[headers_row].height = 36
        static_cols = ['№', 'ФИО', 'Должность']
        for ci, h in enumerate(static_cols, 1):
            c = ws.cell(row=headers_row, column=ci, value=h)
            c.font = Font(name='Times New Roman', size=10, bold=True)
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.border = brd
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        for d in range(1, 32):
            c = ws.cell(row=headers_row, column=3 + d, value=d)
            c.font = Font(name='Times New Roman', size=9, bold=True)
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.border = brd
            c.alignment = Alignment(horizontal='center', vertical='center')

        for c_i, txt in enumerate(['Явки', 'Не явки'], 35):
            c = ws.cell(row=headers_row, column=c_i, value=txt)
            c.font = Font(name='Times New Roman', size=9, bold=True)
            c.fill = PatternFill('solid', fgColor='FEF9C3')
            c.border = brd
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Строки сотрудников
        for i in range(1, 11):
            r = headers_row + i
            ws.row_dimensions[r].height = 18
            ws.cell(row=r, column=1, value=i).border = brd
            for ci, val in enumerate([f'{{{{фио_{i}}}}}', f'{{{{должность_{i}}}}}'], 2):
                c = ws.cell(row=r, column=ci, value=val)
                c.font = Font(name='Times New Roman', size=10)
                c.border = brd
                c.alignment = Alignment(horizontal='left', vertical='center')
            for d in range(1, 32):
                c = ws.cell(row=r, column=3 + d)
                c.border = brd
                c.alignment = Alignment(horizontal='center', vertical='center')

        r_sign = headers_row + 12
        ws[f'A{r_sign}'].value = 'Руководитель:'
        ws[f'A{r_sign}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_sign}'].value = '{{руководитель}}'
        ws[f'C{r_sign}'].font = Font(name='Times New Roman', size=11)
        ws[f'G{r_sign}'].value = '_________________________'

        r_sign2 = r_sign + 1
        ws[f'A{r_sign2}'].value = 'Ответственный:'
        ws[f'A{r_sign2}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_sign2}'].value = '{{ответственный}}'
        ws[f'C{r_sign2}'].font = Font(name='Times New Roman', size=11)

        wb.save(path)
        return path

    def _make_xlsx_avansoviy(self, path):
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = 'Авансовый отчёт'

        thin = Side(style='thin')
        brd = Border(left=thin, right=thin, top=thin, bottom=thin)

        col_w = {1: 6, 2: 12, 3: 35, 4: 14, 5: 14, 6: 14}
        for c, w in col_w.items():
            ws.column_dimensions[get_column_letter(c)].width = w

        # Шапка
        ws.merge_cells('A1:F1')
        ws['A1'].value = '{{компания}}'
        ws['A1'].font = Font(name='Times New Roman', size=13, bold=True)
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[1].height = 22

        ws.merge_cells('A2:F2')
        ws['A2'].value = 'АВАНСОВЫЙ ОТЧЁТ № {{номер_отчёта}} от {{дата}} г.'
        ws['A2'].font = Font(name='Times New Roman', size=14, bold=True, color='FFFFFF')
        ws['A2'].fill = PatternFill('solid', fgColor='2563EB')
        ws['A2'].alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[2].height = 26

        info = [
            ('Подотчётное лицо:', '{{фио_подотчётного}}'),
            ('Должность:', '{{должность}}'),
            ('Подразделение:', '{{отдел}}'),
            ('Назначение аванса:', '{{назначение}}'),
            ('Командировка / период:', '{{даты_командировки}}'),
        ]
        for i, (lbl, val) in enumerate(info, 3):
            ws.merge_cells(f'A{i}:B{i}')
            ws[f'A{i}'].value = lbl
            ws[f'A{i}'].font = Font(name='Times New Roman', size=11, bold=True)
            ws.merge_cells(f'C{i}:F{i}')
            ws[f'C{i}'].value = val
            ws[f'C{i}'].font = Font(name='Times New Roman', size=11)
            ws.row_dimensions[i].height = 18

        ws.row_dimensions[8].height = 6

        # Заголовки таблицы расходов
        hr = 9
        ws.row_dimensions[hr].height = 32
        for ci, (txt, cen) in enumerate([
            ('№', True), ('Дата', True),
            ('Наименование расхода / документ', False),
            ('Сумма (руб.)', True), ('НДС', True), ('Принято', True),
        ], 1):
            c = ws.cell(row=hr, column=ci, value=txt)
            c.font = Font(name='Times New Roman', size=11, bold=True)
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.border = brd
            c.alignment = Alignment(horizontal='center' if cen else 'left', vertical='center', wrap_text=True)

        for i in range(1, 9):
            r = hr + i
            ws.row_dimensions[r].height = 20
            for ci, val in enumerate([
                str(i),
                f'{{{{дата_расхода_{i}}}}}',
                f'{{{{наименование_расхода_{i}}}}}',
                f'{{{{сумма_расхода_{i}}}}}',
                f'{{{{ндс_расхода_{i}}}}}',
                f'{{{{принято_{i}}}}}',
            ], 1):
                c = ws.cell(row=r, column=ci, value=val)
                c.font = Font(name='Times New Roman', size=11)
                c.border = brd
                c.alignment = Alignment(horizontal='center' if ci != 3 else 'left', vertical='center')

        r_tot = hr + 9
        ws.merge_cells(f'A{r_tot}:C{r_tot}')
        ws[f'A{r_tot}'].value = 'ИТОГО РАСХОДОВ:'
        ws[f'A{r_tot}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'A{r_tot}'].alignment = Alignment(horizontal='right', vertical='center')
        ws[f'A{r_tot}'].fill = PatternFill('solid', fgColor='DBEAFE')
        for ci, val in [(4, '{{итого_расходов}}'), (5, '{{итого_ндс}}'), (6, '{{итого_принято}}')]:
            c = ws.cell(row=r_tot, column=ci, value=val)
            c.font = Font(name='Times New Roman', size=11, bold=True)
            c.border = brd
            c.fill = PatternFill('solid', fgColor='DBEAFE')
            c.alignment = Alignment(horizontal='center', vertical='center')
        ws.row_dimensions[r_tot].height = 22

        r_s = r_tot + 2
        ws[f'A{r_s}'].value = 'Выдан аванс:'
        ws[f'A{r_s}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_s}'].value = '{{выдан_аванс}}'
        ws[f'C{r_s}'].font = Font(name='Times New Roman', size=11)

        r_s += 1
        ws[f'A{r_s}'].value = 'Остаток / перерасход:'
        ws[f'A{r_s}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_s}'].value = '{{остаток_перерасход}}'
        ws[f'C{r_s}'].font = Font(name='Times New Roman', size=11)

        r_s += 2
        ws[f'A{r_s}'].value = 'Подотчётное лицо:'
        ws[f'A{r_s}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_s}'].value = '{{фио_подотчётного}}'
        ws[f'C{r_s}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r_s}'].value = '_________________________'

        r_s += 1
        ws[f'A{r_s}'].value = 'Бухгалтер:'
        ws[f'A{r_s}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_s}'].value = '{{бухгалтер}}'
        ws[f'C{r_s}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r_s}'].value = '_________________________'

        r_s += 1
        ws[f'A{r_s}'].value = 'Утверждаю (руководитель):'
        ws[f'A{r_s}'].font = Font(name='Times New Roman', size=11, bold=True)
        ws[f'C{r_s}'].value = '{{руководитель}}'
        ws[f'C{r_s}'].font = Font(name='Times New Roman', size=11)
        ws[f'E{r_s}'].value = '_________________________'

        wb.save(path)
        return path

    # ─── PDF ─────────────────────────────────────────────────────────────────
    def _make_pdf_obyasnitelnaya(self, path):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_paths = [
            'C:/Windows/Fonts/times.ttf',
            'C:/Windows/Fonts/Times New Roman.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
        ]
        font_name = 'Helvetica'
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('TimesOb', fp))
                    font_name = 'TimesOb'
                    break
                except Exception:
                    pass

        c = Canvas(path, pagesize=A4)
        w, h = A4

        def text(x, y, s, size=12):
            c.setFont(font_name, size)
            c.drawString(x, y, s)

        y = h - 2 * cm
        text(w / 2 - 5 * cm, y, 'ОБЪЯСНИТЕЛЬНАЯ ЗАПИСКА', size=15)
        y -= 0.7 * cm
        text(w / 2 - 4 * cm, y, '{{дата}} г.', size=12)
        y -= 1.0 * cm
        text(2 * cm, y, 'Кому: {{кому_должность}} {{кому_фио}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'От: {{фамилия}} {{имя}} {{отчество}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'Должность: {{должность}}', size=12)
        y -= 0.6 * cm
        text(2 * cm, y, 'Подразделение: {{отдел}}', size=12)
        y -= 1.0 * cm

        lines = [
            'Настоящим разъясняю обстоятельства, связанные с {{предмет_объяснения}}:',
            '',
            '{{текст_объяснения}}',
            '',
            'Причиной произошедшего является: {{причина}}.',
            '',
            'Для предотвращения подобных ситуаций в будущем обязуюсь: {{обязательство}}.',
            '',
            '',
            'Приложения: {{приложения}}.',
            '',
            '',
            'Дата: {{дата}}',
            '',
            'Подпись: _______________________  {{фамилия}} {{имя}} {{отчество}}',
        ]
        for line in lines:
            text(2 * cm, y, line, size=12)
            y -= 0.6 * cm

        c.save()
        return path

    def _make_pdf_zayavlenie_priem(self, path):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        font_paths = [
            'C:/Windows/Fonts/times.ttf',
            'C:/Windows/Fonts/Times New Roman.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf',
        ]
        font_name = 'Helvetica'
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('TimesPr', fp))
                    font_name = 'TimesPr'
                    break
                except Exception:
                    pass

        c = Canvas(path, pagesize=A4)
        w, h = A4

        def text(x, y, s, size=12):
            c.setFont(font_name, size)
            c.drawString(x, y, s)

        y = h - 2 * cm
        # Шапка «кому»
        text(w - 10 * cm, y,        '{{должность_руководителя}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        '{{компания}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        '{{руководитель}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        'от {{фамилия}} {{имя}} {{отчество}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        'паспорт: {{серия_паспорта}} № {{номер_паспорта}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        'выдан: {{кем_выдан}}, {{дата_выдачи}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        'адрес: {{адрес}}', size=11)
        y -= 0.55 * cm
        text(w - 10 * cm, y,        'тел.: {{телефон}}', size=11)
        y -= 1.2 * cm

        text(w / 2 - 3.5 * cm, y, 'ЗАЯВЛЕНИЕ', size=15)
        y -= 0.8 * cm
        text(w / 2 - 4.5 * cm, y, 'о приёме на работу', size=13)
        y -= 1.0 * cm

        lines = [
            'Прошу принять меня на работу на должность {{должность}}',
            'в {{отдел}} с {{дата_начала}} г.',
            '',
            'Форма занятости: {{тип_занятости}}.',
            'Желаемый оклад: {{оклад}} рублей.',
            '',
            'Сведения об образовании: {{образование}}.',
            'Специальность: {{специальность}}.',
            'Опыт работы: {{опыт_работы}}.',
            '',
            '',
            'Дата: {{дата}}',
            '',
            'Подпись: _______________________  {{фамилия}} {{имя}} {{отчество}}',
        ]
        for line in lines:
            text(2 * cm, y, line, size=12)
            y -= 0.6 * cm

        c.save()
        return path

    # ══════════════════════════════════════ DB ═══════════════════════════════
    def _create_db_record(self, name, doc_type, file_format, file_path, description, placeholders):
        from documents.models import DocumentTemplate

        # Путь относительно MEDIA_ROOT
        rel = os.path.relpath(file_path, str(settings.MEDIA_ROOT))

        obj, created = DocumentTemplate.objects.update_or_create(
            name=name,
            defaults=dict(
                type=doc_type,
                file_format=file_format,
                description=description,
                placeholders=placeholders,
                is_active=True,
            )
        )
        obj.template_file.name = rel
        obj.save(update_fields=['template_file'])
        return obj, created

    # ══════════════════════════════════════ MAIN ═════════════════════════════
    def handle(self, *args, **options):
        self.stdout.write(self.style.MIGRATE_HEADING('Создание файлов шаблонов...'))

        templates_to_create = [
            # ── DOCX ──────────────────────────────────────────────────────
            {
                'name': 'Приказ о приёме на работу',
                'type': 'order',
                'format': 'docx',
                'builder': self._make_docx_prikaz,
                'filename': 'prikaz_priem.docx',
                'description': 'Приказ о приёме сотрудника на работу (ст. 68 ТК РФ). '
                               'Содержит сведения о должности, окладе, испытательном сроке.',
                'placeholders': [
                    {'name': 'компания',           'label': 'Название организации',    'type': 'text',  'required': True},
                    {'name': 'номер_приказа',       'label': 'Номер приказа',           'type': 'text',  'required': True},
                    {'name': 'дата',                'label': 'Дата приказа',            'type': 'date',  'required': True},
                    {'name': 'номер_договора',      'label': 'Номер трудового договора','type': 'text',  'required': True},
                    {'name': 'дата_договора',       'label': 'Дата трудового договора', 'type': 'date',  'required': True},
                    {'name': 'фамилия',             'label': 'Фамилия',                 'type': 'text',  'required': True},
                    {'name': 'имя',                 'label': 'Имя',                     'type': 'text',  'required': True},
                    {'name': 'отчество',            'label': 'Отчество',                'type': 'text',  'required': False},
                    {'name': 'должность',           'label': 'Должность',               'type': 'text',  'required': True},
                    {'name': 'отдел',               'label': 'Отдел / подразделение',   'type': 'text',  'required': True},
                    {'name': 'дата_начала',         'label': 'Дата начала работы',      'type': 'date',  'required': True},
                    {'name': 'оклад',               'label': 'Оклад (цифрами, руб.)',   'type': 'text',  'required': True},
                    {'name': 'испытательный_срок',  'label': 'Испытательный срок (мес.)','type': 'number','required': True},
                    {'name': 'руководитель_отдела', 'label': 'ФИО руководителя отдела', 'type': 'text',  'required': True},
                    {'name': 'руководитель',        'label': 'ФИО директора',           'type': 'text',  'required': True},
                    {'name': 'сотрудник_подпись',   'label': 'ФИО сотрудника (для подписи)','type': 'text','required': False},
                ],
            },
            {
                'name': 'Трудовой договор',
                'type': 'contract',
                'format': 'docx',
                'builder': self._make_docx_dogovor,
                'filename': 'trudovoy_dogovor.docx',
                'description': 'Стандартный трудовой договор с разделами: предмет, срок, '
                               'оплата труда, режим работы, права и обязанности сторон.',
                'placeholders': [
                    {'name': 'компания',              'label': 'Название организации',       'type': 'text', 'required': True},
                    {'name': 'номер_договора',         'label': 'Номер договора',              'type': 'text', 'required': True},
                    {'name': 'город',                  'label': 'Город',                       'type': 'text', 'required': True},
                    {'name': 'год',                    'label': 'Год',                         'type': 'text', 'required': True},
                    {'name': 'должность_руководителя', 'label': 'Должность подписанта',        'type': 'text', 'required': True},
                    {'name': 'руководитель',           'label': 'ФИО руководителя',            'type': 'text', 'required': True},
                    {'name': 'основание',              'label': 'Основание полномочий',         'type': 'text', 'required': True},
                    {'name': 'фамилия',                'label': 'Фамилия сотрудника',          'type': 'text', 'required': True},
                    {'name': 'имя',                    'label': 'Имя сотрудника',              'type': 'text', 'required': True},
                    {'name': 'отчество',               'label': 'Отчество сотрудника',         'type': 'text', 'required': False},
                    {'name': 'должность',              'label': 'Должность сотрудника',        'type': 'text', 'required': True},
                    {'name': 'отдел',                  'label': 'Отдел / подразделение',       'type': 'text', 'required': True},
                    {'name': 'адрес_работы',           'label': 'Адрес места работы',          'type': 'text', 'required': True},
                    {'name': 'тип_занятости',          'label': 'Тип занятости',               'type': 'text', 'required': True},
                    {'name': 'срок_договора',          'label': 'Срок договора',               'type': 'text', 'required': True},
                    {'name': 'дата_начала',            'label': 'Дата начала работы',          'type': 'date', 'required': True},
                    {'name': 'испытательный_срок',     'label': 'Испытательный срок (мес.)',   'type': 'text', 'required': True},
                    {'name': 'оклад',                  'label': 'Оклад (цифрами)',             'type': 'text', 'required': True},
                    {'name': 'оклад_прописью',         'label': 'Оклад (прописью)',            'type': 'text', 'required': True},
                    {'name': 'дата_выплаты_1',         'label': 'Дата выплаты 1 (число)',      'type': 'text', 'required': True},
                    {'name': 'дата_выплаты_2',         'label': 'Дата выплаты 2 (число)',      'type': 'text', 'required': True},
                    {'name': 'режим_работы',           'label': 'Режим работы (5/2 и т.п.)',   'type': 'text', 'required': True},
                    {'name': 'часы_в_неделю',          'label': 'Часов в неделю',              'type': 'number','required': True},
                    {'name': 'время_начала',           'label': 'Время начала работы',         'type': 'text', 'required': True},
                    {'name': 'время_окончания',        'label': 'Время окончания работы',      'type': 'text', 'required': True},
                    {'name': 'обеденный_перерыв',      'label': 'Перерыв (минут)',             'type': 'number','required': True},
                    {'name': 'отпуск',                 'label': 'Отпуск (дней)',               'type': 'number','required': True},
                    {'name': 'инн_компании',           'label': 'ИНН организации',             'type': 'text', 'required': False},
                    {'name': 'адрес_компании',         'label': 'Адрес организации',           'type': 'text', 'required': False},
                    {'name': 'серия_паспорта',         'label': 'Серия паспорта сотрудника',   'type': 'text', 'required': False},
                    {'name': 'номер_паспорта',         'label': 'Номер паспорта сотрудника',   'type': 'text', 'required': False},
                    {'name': 'адрес_работника',        'label': 'Адрес сотрудника',            'type': 'text', 'required': False},
                ],
            },
            {
                'name': 'Акт выполненных работ',
                'type': 'act',
                'format': 'docx',
                'builder': self._make_docx_akt,
                'filename': 'akt_vypolnennykh_rabot.docx',
                'description': 'Акт приёма-передачи выполненных работ / оказанных услуг '
                               'между двумя организациями. Содержит таблицу работ и подписи.',
                'placeholders': [
                    {'name': 'номер_акта',              'label': 'Номер акта',                  'type': 'text', 'required': True},
                    {'name': 'дата',                    'label': 'Дата акта',                   'type': 'date', 'required': True},
                    {'name': 'компания_заказчик',       'label': 'Название заказчика',          'type': 'text', 'required': True},
                    {'name': 'должность_заказчика',     'label': 'Должность подписанта (заказчик)', 'type': 'text', 'required': True},
                    {'name': 'руководитель_заказчика',  'label': 'ФИО подписанта (заказчик)',   'type': 'text', 'required': True},
                    {'name': 'основание_заказчика',     'label': 'Основание полномочий (заказчик)', 'type': 'text', 'required': True},
                    {'name': 'компания_исполнитель',    'label': 'Название исполнителя',        'type': 'text', 'required': True},
                    {'name': 'должность_исполнителя',   'label': 'Должность подписанта (исполнитель)', 'type': 'text', 'required': True},
                    {'name': 'руководитель_исполнителя','label': 'ФИО подписанта (исполнитель)','type': 'text', 'required': True},
                    {'name': 'основание_исполнителя',   'label': 'Основание полномочий (исполнитель)', 'type': 'text', 'required': True},
                    {'name': 'наименование_работы_1',   'label': 'Работа 1: наименование',      'type': 'text', 'required': True},
                    {'name': 'единица_1',               'label': 'Работа 1: ед. изм.',          'type': 'text', 'required': True},
                    {'name': 'количество_1',            'label': 'Работа 1: количество',        'type': 'text', 'required': True},
                    {'name': 'сумма_1',                 'label': 'Работа 1: сумма (руб.)',      'type': 'text', 'required': True},
                    {'name': 'наименование_работы_2',   'label': 'Работа 2: наименование',      'type': 'text', 'required': False},
                    {'name': 'единица_2',               'label': 'Работа 2: ед. изм.',          'type': 'text', 'required': False},
                    {'name': 'количество_2',            'label': 'Работа 2: количество',        'type': 'text', 'required': False},
                    {'name': 'сумма_2',                 'label': 'Работа 2: сумма (руб.)',      'type': 'text', 'required': False},
                    {'name': 'итого',                   'label': 'Итого сумма (руб.)',          'type': 'text', 'required': True},
                    {'name': 'итого_прописью',          'label': 'Итого (прописью)',            'type': 'text', 'required': True},
                    {'name': 'ндс',                     'label': 'Ставка НДС (%)',              'type': 'text', 'required': True},
                    {'name': 'сумма_ндс',               'label': 'Сумма НДС (руб.)',            'type': 'text', 'required': True},
                ],
            },
            # ── XLSX ──────────────────────────────────────────────────────
            {
                'name': 'Счёт на оплату',
                'type': 'other',
                'format': 'xlsx',
                'builder': self._make_xlsx_schet,
                'filename': 'schet_na_oplatu.xlsx',
                'description': 'Счёт на оплату товаров/услуг с реквизитами поставщика, '
                               'таблицей позиций, расчётом НДС и итоговой суммой.',
                'placeholders': [
                    {'name': 'номер_счета',      'label': 'Номер счёта',              'type': 'text', 'required': True},
                    {'name': 'дата',             'label': 'Дата выставления',         'type': 'date', 'required': True},
                    {'name': 'компания_поставщик','label': 'Название поставщика',     'type': 'text', 'required': True},
                    {'name': 'инн_поставщик',    'label': 'ИНН поставщика',           'type': 'text', 'required': True},
                    {'name': 'кпп_поставщик',    'label': 'КПП поставщика',           'type': 'text', 'required': False},
                    {'name': 'банк_поставщик',   'label': 'Банк поставщика',          'type': 'text', 'required': True},
                    {'name': 'расчетный_счет',   'label': 'Расчётный счёт',           'type': 'text', 'required': True},
                    {'name': 'бик',              'label': 'БИК',                      'type': 'text', 'required': True},
                    {'name': 'компания_покупатель','label': 'Название покупателя',    'type': 'text', 'required': True},
                    {'name': 'адрес_покупателя', 'label': 'Адрес покупателя',         'type': 'text', 'required': True},
                    {'name': 'наименование_1',   'label': 'Позиция 1: наименование',  'type': 'text', 'required': True},
                    {'name': 'единица_1',        'label': 'Позиция 1: ед. изм.',      'type': 'text', 'required': True},
                    {'name': 'количество_1',     'label': 'Позиция 1: количество',    'type': 'text', 'required': True},
                    {'name': 'сумма_1',          'label': 'Позиция 1: сумма',         'type': 'text', 'required': True},
                    {'name': 'ставка_ндс',       'label': 'Ставка НДС (%)',           'type': 'text', 'required': True},
                    {'name': 'сумма_без_ндс',    'label': 'Сумма без НДС',            'type': 'text', 'required': True},
                    {'name': 'сумма_ндс',        'label': 'Сумма НДС',                'type': 'text', 'required': True},
                    {'name': 'итого',            'label': 'Итого к оплате',           'type': 'text', 'required': True},
                    {'name': 'итого_прописью',   'label': 'Итого прописью',           'type': 'text', 'required': True},
                    {'name': 'копейки',          'label': 'Копейки (00)',             'type': 'text', 'required': True},
                    {'name': 'срок_оплаты',      'label': 'Срок оплаты',             'type': 'text', 'required': True},
                    {'name': 'руководитель',     'label': 'ФИО руководителя',         'type': 'text', 'required': True},
                    {'name': 'главный_бухгалтер','label': 'ФИО главного бухгалтера', 'type': 'text', 'required': False},
                ],
            },
            {
                'name': 'Расчётная ведомость по зарплате',
                'type': 'report',
                'format': 'xlsx',
                'builder': self._make_xlsx_vedomost,
                'filename': 'raschet_vedomost.xlsx',
                'description': 'Расчётная ведомость по начислению и выплате заработной платы '
                               'сотрудникам подразделения за выбранный расчётный период.',
                'placeholders': [
                    {'name': 'компания',           'label': 'Название организации',   'type': 'text',  'required': True},
                    {'name': 'месяц',              'label': 'Расчётный месяц',        'type': 'text',  'required': True},
                    {'name': 'год',                'label': 'Год',                    'type': 'text',  'required': True},
                    {'name': 'отдел',              'label': 'Подразделение',          'type': 'text',  'required': True},
                    {'name': 'дата',               'label': 'Дата составления',       'type': 'date',  'required': True},
                    {'name': 'руководитель',       'label': 'ФИО руководителя',       'type': 'text',  'required': True},
                    {'name': 'главный_бухгалтер',  'label': 'ФИО главбуха',          'type': 'text',  'required': False},
                    {'name': 'итого_к_выплате',    'label': 'Итого к выплате',        'type': 'text',  'required': False},
                ],
            },
            # ── PDF ───────────────────────────────────────────────────────
            {
                'name': 'Заявление на отпуск',
                'type': 'application',
                'format': 'pdf',
                'builder': self._make_pdf_zayavlenie,
                'filename': 'zayavlenie_otpusk.pdf',
                'description': 'Заявление сотрудника на предоставление ежегодного '
                               'оплачиваемого отпуска с указанием дат и подписью руководителя.',
                'placeholders': [
                    {'name': 'фамилия',                'label': 'Фамилия сотрудника',            'type': 'text', 'required': True},
                    {'name': 'имя',                    'label': 'Имя сотрудника',                'type': 'text', 'required': True},
                    {'name': 'отчество',               'label': 'Отчество сотрудника',           'type': 'text', 'required': False},
                    {'name': 'должность',              'label': 'Должность сотрудника',          'type': 'text', 'required': True},
                    {'name': 'отдел',                  'label': 'Подразделение',                 'type': 'text', 'required': True},
                    {'name': 'компания',               'label': 'Название организации',          'type': 'text', 'required': True},
                    {'name': 'должность_руководителя', 'label': 'Должность руководителя (кому)', 'type': 'text', 'required': True},
                    {'name': 'руководитель',           'label': 'ФИО руководителя',             'type': 'text', 'required': True},
                    {'name': 'продолжительность_отпуска','label': 'Дней отпуска',               'type': 'number','required': True},
                    {'name': 'дата_начала',            'label': 'Дата начала отпуска',          'type': 'date', 'required': True},
                    {'name': 'дата_окончания',         'label': 'Дата окончания отпуска',       'type': 'date', 'required': True},
                    {'name': 'год',                    'label': 'Год (отпускной)',              'type': 'text', 'required': True},
                    {'name': 'дата',                   'label': 'Дата написания заявления',     'type': 'date', 'required': True},
                    {'name': 'дата_согласования',      'label': 'Дата согласования руководителем','type': 'date','required': False},
                ],
            },
            {
                'name': 'Служебная записка',
                'type': 'memo',
                'format': 'pdf',
                'builder': self._make_pdf_memo,
                'filename': 'sluzhebnaya_zapiska.pdf',
                'description': 'Внутренний документ для передачи информации или запроса '
                               'между подразделениями. Содержит поля кому, от кого, тема, текст, просьба.',
                'placeholders': [
                    {'name': 'номер',            'label': 'Номер записки',            'type': 'text',    'required': True},
                    {'name': 'дата',             'label': 'Дата',                     'type': 'date',    'required': True},
                    {'name': 'кому_должность',   'label': 'Кому: должность',          'type': 'text',    'required': True},
                    {'name': 'кому_фио',         'label': 'Кому: ФИО',               'type': 'text',    'required': True},
                    {'name': 'от_должность',     'label': 'От кого: должность',       'type': 'text',    'required': True},
                    {'name': 'от_фио',           'label': 'От кого: ФИО',            'type': 'text',    'required': True},
                    {'name': 'тема',             'label': 'Тема / заголовок',         'type': 'text',    'required': True},
                    {'name': 'текст_сообщения',  'label': 'Текст сообщения',          'type': 'textarea','required': True},
                    {'name': 'просьба',          'label': 'Просьба / запрос',         'type': 'textarea','required': True},
                    {'name': 'приложения',       'label': 'Приложения (или "нет")',   'type': 'text',    'required': False},
                ],
            },
            # ── DOCX (new) ────────────────────────────────────────────────
            {
                'name': 'Доверенность',
                'type': 'order',
                'format': 'docx',
                'builder': self._make_docx_doverennost,
                'filename': 'doverennost.docx',
                'description': 'Доверенность на представление интересов организации. '
                               'Содержит данные доверителя, поверенного, цель и срок действия.',
                'placeholders': [
                    {'name': 'компания',              'label': 'Название организации',        'type': 'text',  'required': True},
                    {'name': 'номер_доверенности',     'label': 'Номер доверенности',          'type': 'text',  'required': True},
                    {'name': 'город',                  'label': 'Город',                       'type': 'text',  'required': True},
                    {'name': 'дата',                   'label': 'Дата доверенности',           'type': 'date',  'required': True},
                    {'name': 'должность_руководителя', 'label': 'Должность подписанта',        'type': 'text',  'required': True},
                    {'name': 'руководитель',           'label': 'ФИО руководителя',            'type': 'text',  'required': True},
                    {'name': 'основание',              'label': 'Основание полномочий',         'type': 'text',  'required': True},
                    {'name': 'фамилия',                'label': 'Фамилия поверенного',         'type': 'text',  'required': True},
                    {'name': 'имя',                    'label': 'Имя поверенного',             'type': 'text',  'required': True},
                    {'name': 'отчество',               'label': 'Отчество поверенного',        'type': 'text',  'required': False},
                    {'name': 'серия_паспорта',         'label': 'Серия паспорта поверенного',  'type': 'text',  'required': True},
                    {'name': 'номер_паспорта',         'label': 'Номер паспорта поверенного',  'type': 'text',  'required': True},
                    {'name': 'кем_выдан',              'label': 'Кем выдан паспорт',           'type': 'text',  'required': True},
                    {'name': 'дата_выдачи',            'label': 'Дата выдачи паспорта',        'type': 'date',  'required': True},
                    {'name': 'адрес',                  'label': 'Адрес поверенного',           'type': 'text',  'required': True},
                    {'name': 'место_представления',    'label': 'Место представления',         'type': 'text',  'required': True},
                    {'name': 'предмет_доверенности',   'label': 'Предмет / цель доверенности', 'type': 'text',  'required': True},
                    {'name': 'перечень_действий',      'label': 'Перечень действий',           'type': 'textarea','required': True},
                    {'name': 'срок_действия',          'label': 'Срок действия (дата)',        'type': 'date',  'required': True},
                ],
            },
            {
                'name': 'Договор купли-продажи',
                'type': 'contract',
                'format': 'docx',
                'builder': self._make_docx_kuplia_prodazha,
                'filename': 'dogovor_kupli_prodazhi.docx',
                'description': 'Договор купли-продажи товара между двумя сторонами. '
                               'Содержит предмет, цену, сроки передачи, ответственность.',
                'placeholders': [
                    {'name': 'номер_договора',    'label': 'Номер договора',             'type': 'text',  'required': True},
                    {'name': 'город',             'label': 'Город',                      'type': 'text',  'required': True},
                    {'name': 'дата',              'label': 'Дата договора',              'type': 'date',  'required': True},
                    {'name': 'продавец',          'label': 'Название / ФИО продавца',    'type': 'text',  'required': True},
                    {'name': 'покупатель',        'label': 'Название / ФИО покупателя',  'type': 'text',  'required': True},
                    {'name': 'наименование_товара','label': 'Наименование товара',        'type': 'text',  'required': True},
                    {'name': 'количество',        'label': 'Количество',                 'type': 'text',  'required': True},
                    {'name': 'единица_измерения', 'label': 'Единица измерения',          'type': 'text',  'required': True},
                    {'name': 'характеристики',    'label': 'Характеристики товара',      'type': 'textarea','required': False},
                    {'name': 'адрес_передачи',    'label': 'Адрес передачи товара',      'type': 'text',  'required': True},
                    {'name': 'цена_единицы',      'label': 'Цена единицы (руб.)',        'type': 'text',  'required': True},
                    {'name': 'общая_стоимость',   'label': 'Общая стоимость (руб.)',     'type': 'text',  'required': True},
                    {'name': 'стоимость_прописью','label': 'Стоимость прописью',         'type': 'text',  'required': True},
                    {'name': 'ндс',               'label': 'Ставка НДС (%)',             'type': 'text',  'required': True},
                    {'name': 'срок_оплаты',       'label': 'Срок оплаты (дней)',         'type': 'text',  'required': True},
                    {'name': 'дата_передачи',     'label': 'Дата передачи товара',       'type': 'date',  'required': True},
                    {'name': 'пеня',              'label': 'Пеня за просрочку (%)',      'type': 'text',  'required': False},
                    {'name': 'инн_продавца',      'label': 'ИНН продавца',               'type': 'text',  'required': False},
                    {'name': 'инн_покупателя',    'label': 'ИНН покупателя',             'type': 'text',  'required': False},
                    {'name': 'счет_продавца',     'label': 'Р/с продавца',               'type': 'text',  'required': False},
                    {'name': 'счет_покупателя',   'label': 'Р/с покупателя',             'type': 'text',  'required': False},
                ],
            },
            {
                'name': 'Приказ об увольнении',
                'type': 'order',
                'format': 'docx',
                'builder': self._make_docx_uvolnenie,
                'filename': 'prikaz_uvolneniye.docx',
                'description': 'Приказ о прекращении (расторжении) трудового договора с работником. '
                               'Содержит основание, дату, реквизиты сотрудника.',
                'placeholders': [
                    {'name': 'компания',            'label': 'Название организации',       'type': 'text',  'required': True},
                    {'name': 'номер_приказа',        'label': 'Номер приказа',              'type': 'text',  'required': True},
                    {'name': 'дата',                 'label': 'Дата приказа',               'type': 'date',  'required': True},
                    {'name': 'основание_увольнения', 'label': 'Основание увольнения',       'type': 'text',  'required': True},
                    {'name': 'статья_тк',            'label': 'Статья ТК РФ',               'type': 'text',  'required': True},
                    {'name': 'номер_договора',       'label': 'Номер трудового договора',   'type': 'text',  'required': True},
                    {'name': 'дата_договора',        'label': 'Дата трудового договора',    'type': 'date',  'required': True},
                    {'name': 'дата_увольнения',      'label': 'Дата увольнения',            'type': 'date',  'required': True},
                    {'name': 'фамилия',              'label': 'Фамилия',                    'type': 'text',  'required': True},
                    {'name': 'имя',                  'label': 'Имя',                        'type': 'text',  'required': True},
                    {'name': 'отчество',             'label': 'Отчество',                   'type': 'text',  'required': False},
                    {'name': 'табельный_номер',      'label': 'Табельный номер',             'type': 'text',  'required': False},
                    {'name': 'отдел',                'label': 'Подразделение',               'type': 'text',  'required': True},
                    {'name': 'должность',            'label': 'Должность',                  'type': 'text',  'required': True},
                    {'name': 'документ_основание',   'label': 'Документ-основание',         'type': 'text',  'required': True},
                    {'name': 'дата_документа',       'label': 'Дата документа-основания',   'type': 'date',  'required': True},
                    {'name': 'руководитель',         'label': 'ФИО директора',              'type': 'text',  'required': True},
                    {'name': 'дата_ознакомления',    'label': 'Дата ознакомления работника', 'type': 'date',  'required': False},
                ],
            },
            {
                'name': 'Протокол совещания',
                'type': 'memo',
                'format': 'docx',
                'builder': self._make_docx_protokol,
                'filename': 'protokol_soveshania.docx',
                'description': 'Протокол проведения рабочего совещания. Содержит повестку дня, '
                               'участников, содержание докладов, принятые решения и ответственных.',
                'placeholders': [
                    {'name': 'компания',        'label': 'Название организации',  'type': 'text',  'required': True},
                    {'name': 'номер_протокола', 'label': 'Номер протокола',        'type': 'text',  'required': True},
                    {'name': 'дата',            'label': 'Дата совещания',         'type': 'date',  'required': True},
                    {'name': 'место',           'label': 'Место проведения',       'type': 'text',  'required': True},
                    {'name': 'время_начала',    'label': 'Время начала',           'type': 'text',  'required': True},
                    {'name': 'время_окончания', 'label': 'Время окончания',        'type': 'text',  'required': True},
                    {'name': 'председатель',    'label': 'Председатель (ФИО)',     'type': 'text',  'required': True},
                    {'name': 'секретарь',       'label': 'Секретарь (ФИО)',        'type': 'text',  'required': True},
                    {'name': 'участник_1',      'label': 'Участник 1 (ФИО)',       'type': 'text',  'required': True},
                    {'name': 'должность_участника_1','label': 'Должность участника 1','type': 'text','required': True},
                    {'name': 'участник_2',      'label': 'Участник 2 (ФИО)',       'type': 'text',  'required': False},
                    {'name': 'должность_участника_2','label': 'Должность участника 2','type': 'text','required': False},
                    {'name': 'вопрос_1',        'label': 'Вопрос 1 повестки',      'type': 'text',  'required': True},
                    {'name': 'докладчик_1',     'label': 'Докладчик 1',            'type': 'text',  'required': True},
                    {'name': 'содержание_доклада_1','label': 'Содержание доклада 1','type': 'textarea','required': True},
                    {'name': 'решение_1',       'label': 'Принятое решение 1',     'type': 'textarea','required': True},
                    {'name': 'ответственный_1', 'label': 'Ответственный 1',        'type': 'text',  'required': True},
                    {'name': 'срок_1',          'label': 'Срок исполнения 1',      'type': 'date',  'required': True},
                    {'name': 'вопрос_2',        'label': 'Вопрос 2 повестки',      'type': 'text',  'required': False},
                    {'name': 'докладчик_2',     'label': 'Докладчик 2',            'type': 'text',  'required': False},
                    {'name': 'содержание_доклада_2','label': 'Содержание доклада 2','type': 'textarea','required': False},
                    {'name': 'решение_2',       'label': 'Принятое решение 2',     'type': 'textarea','required': False},
                    {'name': 'ответственный_2', 'label': 'Ответственный 2',        'type': 'text',  'required': False},
                    {'name': 'срок_2',          'label': 'Срок исполнения 2',      'type': 'date',  'required': False},
                ],
            },
            # ── XLSX (new) ────────────────────────────────────────────────
            {
                'name': 'Табель учёта рабочего времени',
                'type': 'report',
                'format': 'xlsx',
                'builder': self._make_xlsx_tabel,
                'filename': 'tabel_ucheta.xlsx',
                'description': 'Табель учёта рабочего времени сотрудников за месяц. '
                               'Содержит ФИО, должности и сетку по дням месяца (31 день).',
                'placeholders': [
                    {'name': 'компания',       'label': 'Название организации',  'type': 'text',  'required': True},
                    {'name': 'месяц',          'label': 'Расчётный месяц',        'type': 'text',  'required': True},
                    {'name': 'год',            'label': 'Год',                    'type': 'text',  'required': True},
                    {'name': 'отдел',          'label': 'Подразделение',          'type': 'text',  'required': True},
                    {'name': 'руководитель',   'label': 'ФИО руководителя',       'type': 'text',  'required': True},
                    {'name': 'ответственный',  'label': 'ФИО ответственного',     'type': 'text',  'required': True},
                    {'name': 'фио_1',          'label': 'Сотрудник 1: ФИО',       'type': 'text',  'required': True},
                    {'name': 'должность_1',    'label': 'Сотрудник 1: должность', 'type': 'text',  'required': True},
                    {'name': 'фио_2',          'label': 'Сотрудник 2: ФИО',       'type': 'text',  'required': False},
                    {'name': 'должность_2',    'label': 'Сотрудник 2: должность', 'type': 'text',  'required': False},
                    {'name': 'фио_3',          'label': 'Сотрудник 3: ФИО',       'type': 'text',  'required': False},
                    {'name': 'должность_3',    'label': 'Сотрудник 3: должность', 'type': 'text',  'required': False},
                ],
            },
            {
                'name': 'Авансовый отчёт',
                'type': 'report',
                'format': 'xlsx',
                'builder': self._make_xlsx_avansoviy,
                'filename': 'avansoviy_otchet.xlsx',
                'description': 'Авансовый отчёт подотчётного лица по командировочным '
                               'и хозяйственным расходам. До 8 позиций расходов.',
                'placeholders': [
                    {'name': 'компания',             'label': 'Название организации',       'type': 'text',  'required': True},
                    {'name': 'номер_отчёта',         'label': 'Номер авансового отчёта',    'type': 'text',  'required': True},
                    {'name': 'дата',                 'label': 'Дата составления',           'type': 'date',  'required': True},
                    {'name': 'фио_подотчётного',     'label': 'ФИО подотчётного лица',      'type': 'text',  'required': True},
                    {'name': 'должность',            'label': 'Должность',                  'type': 'text',  'required': True},
                    {'name': 'отдел',                'label': 'Подразделение',               'type': 'text',  'required': True},
                    {'name': 'назначение',           'label': 'Назначение аванса',          'type': 'text',  'required': True},
                    {'name': 'даты_командировки',    'label': 'Период командировки',        'type': 'text',  'required': False},
                    {'name': 'дата_расхода_1',       'label': 'Расход 1: дата',             'type': 'date',  'required': True},
                    {'name': 'наименование_расхода_1','label': 'Расход 1: наименование',    'type': 'text',  'required': True},
                    {'name': 'сумма_расхода_1',      'label': 'Расход 1: сумма',            'type': 'text',  'required': True},
                    {'name': 'ндс_расхода_1',        'label': 'Расход 1: НДС',              'type': 'text',  'required': False},
                    {'name': 'принято_1',            'label': 'Расход 1: принято',          'type': 'text',  'required': False},
                    {'name': 'итого_расходов',       'label': 'Итого расходов',             'type': 'text',  'required': True},
                    {'name': 'итого_ндс',            'label': 'Итого НДС',                  'type': 'text',  'required': False},
                    {'name': 'итого_принято',        'label': 'Итого принято',              'type': 'text',  'required': False},
                    {'name': 'выдан_аванс',          'label': 'Выдан аванс (руб.)',         'type': 'text',  'required': True},
                    {'name': 'остаток_перерасход',   'label': 'Остаток / перерасход',       'type': 'text',  'required': False},
                    {'name': 'бухгалтер',            'label': 'ФИО бухгалтера',             'type': 'text',  'required': False},
                    {'name': 'руководитель',         'label': 'ФИО руководителя',           'type': 'text',  'required': True},
                ],
            },
            # ── PDF (new) ─────────────────────────────────────────────────
            {
                'name': 'Объяснительная записка',
                'type': 'memo',
                'format': 'pdf',
                'builder': self._make_pdf_obyasnitelnaya,
                'filename': 'obyasnitelnaya_zapiska.pdf',
                'description': 'Объяснительная записка сотрудника по факту нарушения '
                               'или нештатной ситуации. Содержит причину и обязательства.',
                'placeholders': [
                    {'name': 'дата',               'label': 'Дата',                        'type': 'date',     'required': True},
                    {'name': 'кому_должность',     'label': 'Кому: должность',             'type': 'text',     'required': True},
                    {'name': 'кому_фио',           'label': 'Кому: ФИО',                   'type': 'text',     'required': True},
                    {'name': 'фамилия',            'label': 'Фамилия сотрудника',          'type': 'text',     'required': True},
                    {'name': 'имя',                'label': 'Имя сотрудника',              'type': 'text',     'required': True},
                    {'name': 'отчество',           'label': 'Отчество сотрудника',         'type': 'text',     'required': False},
                    {'name': 'должность',          'label': 'Должность сотрудника',        'type': 'text',     'required': True},
                    {'name': 'отдел',              'label': 'Подразделение',               'type': 'text',     'required': True},
                    {'name': 'предмет_объяснения', 'label': 'Предмет объяснения',          'type': 'text',     'required': True},
                    {'name': 'текст_объяснения',   'label': 'Основной текст',              'type': 'textarea', 'required': True},
                    {'name': 'причина',            'label': 'Причина произошедшего',       'type': 'textarea', 'required': True},
                    {'name': 'обязательство',      'label': 'Обязательство на будущее',    'type': 'textarea', 'required': True},
                    {'name': 'приложения',         'label': 'Приложения (или "нет")',      'type': 'text',     'required': False},
                ],
            },
            {
                'name': 'Заявление о приёме на работу',
                'type': 'application',
                'format': 'pdf',
                'builder': self._make_pdf_zayavlenie_priem,
                'filename': 'zayavlenie_priem.pdf',
                'description': 'Заявление кандидата о приёме на работу с указанием '
                               'желаемой должности, оклада, образования и опыта.',
                'placeholders': [
                    {'name': 'должность_руководителя','label': 'Кому: должность',          'type': 'text',  'required': True},
                    {'name': 'компания',              'label': 'Название организации',      'type': 'text',  'required': True},
                    {'name': 'руководитель',          'label': 'ФИО руководителя',          'type': 'text',  'required': True},
                    {'name': 'фамилия',               'label': 'Фамилия заявителя',         'type': 'text',  'required': True},
                    {'name': 'имя',                   'label': 'Имя заявителя',             'type': 'text',  'required': True},
                    {'name': 'отчество',              'label': 'Отчество заявителя',        'type': 'text',  'required': False},
                    {'name': 'серия_паспорта',        'label': 'Серия паспорта',            'type': 'text',  'required': False},
                    {'name': 'номер_паспорта',        'label': 'Номер паспорта',            'type': 'text',  'required': False},
                    {'name': 'кем_выдан',             'label': 'Кем выдан паспорт',        'type': 'text',  'required': False},
                    {'name': 'дата_выдачи',           'label': 'Дата выдачи паспорта',     'type': 'date',  'required': False},
                    {'name': 'адрес',                 'label': 'Адрес проживания',          'type': 'text',  'required': False},
                    {'name': 'телефон',               'label': 'Телефон',                   'type': 'text',  'required': False},
                    {'name': 'должность',             'label': 'Желаемая должность',        'type': 'text',  'required': True},
                    {'name': 'отдел',                 'label': 'Желаемый отдел',            'type': 'text',  'required': True},
                    {'name': 'дата_начала',           'label': 'Дата начала работы',        'type': 'date',  'required': True},
                    {'name': 'тип_занятости',         'label': 'Форма занятости',           'type': 'text',  'required': True},
                    {'name': 'оклад',                 'label': 'Желаемый оклад (руб.)',     'type': 'text',  'required': False},
                    {'name': 'образование',           'label': 'Образование',               'type': 'text',  'required': False},
                    {'name': 'специальность',         'label': 'Специальность',             'type': 'text',  'required': False},
                    {'name': 'опыт_работы',           'label': 'Опыт работы',               'type': 'textarea','required': False},
                    {'name': 'дата',                  'label': 'Дата подписания',           'type': 'date',  'required': True},
                ],
            },
        ]

        created_count = 0
        updated_count = 0

        for tpl in templates_to_create:
            file_path = self._media_path(tpl['filename'])
            try:
                tpl['builder'](file_path)
                obj, created = self._create_db_record(
                    name=tpl['name'],
                    doc_type=tpl['type'],
                    file_format=tpl['format'],
                    file_path=file_path,
                    description=tpl['description'],
                    placeholders=tpl['placeholders'],
                )
                if created:
                    created_count += 1
                    self.stdout.write(self.style.SUCCESS(
                        f'  ✓ Создан: [{tpl["format"].upper()}] {tpl["name"]}'
                    ))
                else:
                    updated_count += 1
                    self.stdout.write(self.style.WARNING(
                        f'  ↺ Обновлён: [{tpl["format"].upper()}] {tpl["name"]}'
                    ))
            except Exception as e:
                self.stdout.write(self.style.ERROR(
                    f'  ✗ Ошибка при создании "{tpl["name"]}": {e}'
                ))
                import traceback
                traceback.print_exc()

        self.stdout.write('')
        self.stdout.write(self.style.SUCCESS(
            f'Готово! Создано: {created_count}, обновлено: {updated_count} шаблонов.'
        ))
