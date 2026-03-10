"""
Утилиты для работы с Office документами (Word, Excel, PDF)
"""
import os
import re
from io import BytesIO
from datetime import datetime

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


def replace_placeholders_in_text(text, replacements):
    """
    Заменяет плейсхолдеры в тексте
    Поддерживает форматы: {{placeholder}}, {placeholder}, [placeholder]
    """
    if not text:
        return text
    
    result = text
    for key, value in replacements.items():
        # Поддержка разных форматов плейсхолдеров
        patterns = [
            r'\{\{' + re.escape(key) + r'\}\}',  # {{key}}
            r'\{' + re.escape(key) + r'\}',       # {key}
            r'\[' + re.escape(key) + r'\]',       # [key]
        ]
        
        for pattern in patterns:
            result = re.sub(pattern, str(value), result, flags=re.IGNORECASE)
    
    return result


def _replace_in_paragraph(paragraph, replacements):
    """
    Заменяет плейсхолдеры в параграфе, собирая весь текст параграфа в одну строку.
    Word часто дробит плейсхолдер {{name}} по нескольким runs—этот метод решает её.
    """
    if not paragraph.runs:
        return
    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = replace_placeholders_in_text(full_text, replacements)
    if new_text != full_text:
        # Переносим весь текст в первый run (сохраняет его форматирование), остальные очищаем
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def generate_word_document(template_path, output_path, replacements):
    """
    Генерирует Word документ из шаблона с заменой плейсхолдеров
    
    Args:
        template_path: путь к файлу шаблона .docx
        output_path: путь для сохранения результата
        replacements: словарь замен {плейсхолдер: значение}
    
    Returns:
        True если успешно, False если ошибка
    """
    if not DOCX_AVAILABLE:
        raise ImportError("Библиотека python-docx не установлена. Установите: pip install python-docx")
    
    try:
        # Открываем шаблон
        doc = DocxDocument(template_path)
        
        # Заменяем в параграфах
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        
        # Заменяем в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, replacements)
        
        # Заменяем в колонтитулах
        for section in doc.sections:
            # Верхний колонтитул
            header = section.header
            for paragraph in header.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            
            # Нижний колонтитул
            footer = section.footer
            for paragraph in footer.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
        
        # Сохраняем
        doc.save(output_path)
        return True
    
    except Exception as e:
        print(f"Ошибка генерации Word документа: {e}")
        return False


def generate_excel_document(template_path, output_path, replacements):
    """
    Генерирует Excel документ из шаблона с заменой плейсхолдеров
    
    Args:
        template_path: путь к файлу шаблона .xlsx
        output_path: путь для сохранения результата
        replacements: словарь замен {плейсхолдер: значение}
    
    Returns:
        True если успешно, False если ошибка
    """
    if not OPENPYXL_AVAILABLE:
        raise ImportError("Библиотека openpyxl не установлена. Установите: pip install openpyxl")
    
    try:
        # Открываем шаблон
        workbook = load_workbook(template_path)
        
        # Проходим по всем листам
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Проходим по всем ячейкам
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell.value = replace_placeholders_in_text(cell.value, replacements)
        
        # Сохраняем
        workbook.save(output_path)
        return True
    
    except Exception as e:
        print(f"Ошибка генерации Excel документа: {e}")
        return False



def generate_pdf_document(template_path, output_path, replacements):
    """
    Regenerates a ReportLab PDF template with filled placeholder values.
    Works by monkey-patching Canvas.drawString to apply replacements on every text call,
    then re-running the original PDF builder function.
    """
    try:
        from reportlab.pdfgen.canvas import Canvas as _Canvas

        # Map template filename → builder method name in create_template_files.Command
        _BUILDER_MAP = {
            'zayavlenie_otpusk.pdf':      '_make_pdf_zayavlenie',
            'sluzhebnaya_zapiska.pdf':    '_make_pdf_memo',
            'obyasnitelnaya_zapiska.pdf': '_make_pdf_obyasnitelnaya',
            'zayavlenie_priem.pdf':       '_make_pdf_zayavlenie_priem',
        }

        filename = os.path.basename(template_path)
        builder_name = _BUILDER_MAP.get(filename)
        if not builder_name:
            return False, f"Не найден генератор для PDF шаблона: {filename}"

        from documents.management.commands.create_template_files import Command as _Cmd
        cmd = _Cmd()

        # Monkey-patch drawString to apply replacements before drawing
        orig_drawString = _Canvas.drawString

        def _patched_drawString(self, x, y, text, **kwargs):
            replaced = replace_placeholders_in_text(str(text), replacements)
            return orig_drawString(self, x, y, replaced, **kwargs)

        _Canvas.drawString = _patched_drawString
        try:
            getattr(cmd, builder_name)(output_path)
        finally:
            _Canvas.drawString = orig_drawString  # always restore

        return True, None

    except Exception as e:
        print(f"Ошибка генерации PDF документа: {e}")
        return False, str(e)


def generate_document_from_template(template_file_path, template_format, output_path, replacements):
    """
    Универсальная функция генерации документа из шаблона

    Args:
        template_file_path: путь к файлу шаблона
        template_format: формат документа ('docx', 'xlsx', 'pdf')
        output_path: путь для сохранения результата
        replacements: словарь замен {плейсхолдер: значение}

    Returns:
        tuple (success: bool, error_message: str)
    """
    try:
        if template_format == 'docx':
            if not template_file_path or not os.path.exists(template_file_path):
                return False, "Файл шаблона Word не найден"
            success = generate_word_document(template_file_path, output_path, replacements)
            return (True, None) if success else (False, "Ошибка генерации Word документа")

        elif template_format == 'xlsx':
            if not template_file_path or not os.path.exists(template_file_path):
                return False, "Файл шаблона Excel не найден"
            success = generate_excel_document(template_file_path, output_path, replacements)
            return (True, None) if success else (False, "Ошибка генерации Excel документа")

        elif template_format == 'pdf':
            if not template_file_path or not os.path.exists(template_file_path):
                return False, "Файл шаблона PDF не найден"
            return generate_pdf_document(template_file_path, output_path, replacements)

        else:
            return False, f"Неподдерживаемый формат: {template_format}"

    except ImportError as e:
        return False, str(e)
    except Exception as e:
        return False, f"Непредвиденная ошибка: {str(e)}"


def check_office_libraries():
    """
    Проверяет наличие установленных библиотек для работы с Office
    
    Returns:
        dict с информацией о доступности библиотек
    """
    return {
        'docx': DOCX_AVAILABLE,
        'xlsx': OPENPYXL_AVAILABLE,
        'all_available': DOCX_AVAILABLE and OPENPYXL_AVAILABLE
    }


def _build_sig_rows(signatures):
    """Возвращает список dict с данными подписантов (парсит certificate_info JSON)."""
    import json
    rows = []
    for sig in signatures:
        try:
            cert = json.loads(sig.certificate_info) if sig.certificate_info else {}
        except (ValueError, TypeError):
            cert = {}
        subj = cert.get('subject', {})
        rows.append({
            'full_name':    subj.get('cn') or sig.signer.get_full_name() or sig.signer.username,
            'position':     subj.get('position', ''),
            'organization': subj.get('organization', 'ООО «Первый ключ»'),
            'inn':          subj.get('inn', ''),
            'snils':        subj.get('snils', ''),
            'algorithm':    cert.get('algorithm', 'ГОСТ Р 34.10-2012'),
            'serial':       cert.get('serial_number', ''),
            'issuer':       cert.get('issuer', {}).get('cn', ''),
            'valid_from':   cert.get('valid_from', ''),
            'valid_to':     cert.get('valid_to', ''),
            'thumbprint':   cert.get('thumbprint', ''),
            'hash':         sig.signature_data or '',
            'signed_at':    sig.signed_at.strftime('%d.%m.%Y в %H:%M:%S') if sig.signed_at else '',
            'ip':           sig.ip_address or '',
            'is_valid':     sig.is_valid,
        })
    return rows


def append_ep_stamp_to_docx(docx_path, signatures):
    """
    Добавляет блок ЭП в конец существующего .docx файла.

    Args:
        docx_path: абсолютный путь к .docx файлу (будет перезаписан)
        signatures: QuerySet или список ElectronicSignature
    """
    if not DOCX_AVAILABLE:
        return False, 'python-docx не установлен'
    if not os.path.exists(docx_path):
        return False, 'Файл документа не найден'

    try:
        from docx.shared import Pt, RGBColor, Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = DocxDocument(docx_path)
        rows = _build_sig_rows(signatures)
        if not rows:
            return True, None

        # --- разделитель ---
        hr = doc.add_paragraph()
        hr.paragraph_format.space_before = Pt(12)
        run = hr.add_run('─' * 72)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)

        # --- заголовок блока ---
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run('КВАЛИФИЦИРОВАННАЯ ЭЛЕКТРОННАЯ ПОДПИСЬ')
        title_run.bold = True
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = RGBColor(0x06, 0x4e, 0x3b)

        for r in rows:
            status_label = '✔ Действительна' if r['is_valid'] else '✗ Недействительна'
            status_color = RGBColor(0x06, 0x4e, 0x3b) if r['is_valid'] else RGBColor(0x99, 0x1b, 0x1b)

            # таблица 2×N
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            def _row(label, value, bold_val=False):
                tr = tbl.add_row()
                lc, vc = tr.cells[0], tr.cells[1]
                lc.width = Cm(5)
                lr = lc.paragraphs[0].add_run(label)
                lr.font.size = Pt(8)
                lr.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
                vr = vc.paragraphs[0].add_run(str(value))
                vr.font.size = Pt(8.5)
                if bold_val:
                    vr.bold = True

            _row('Подписант', r['full_name'], bold_val=True)
            _row('Статус подписи', status_label)
            if r['position']:
                _row('Должность', r['position'])
            if r['organization']:
                _row('Организация', r['organization'])
            if r['inn']:
                _row('ИНН', r['inn'])
            if r['snils']:
                _row('СНИЛС', r['snils'])
            _row('Дата подписания', r['signed_at'])
            if r['ip']:
                _row('IP адрес', r['ip'])
            _row('Алгоритм', r['algorithm'])
            if r['issuer']:
                _row('Издатель (УЦ)', r['issuer'])
            if r['serial']:
                _row('Серийный номер', r['serial'])
            if r['valid_from'] and r['valid_to']:
                _row('Срок действия', f"{r['valid_from']} — {r['valid_to']}")
            if r['thumbprint']:
                _row('Отпечаток SHA-1', r['thumbprint'])
            _row('Хеш документа (SHA-256)', r['hash'])

            doc.add_paragraph()  # отступ между подписантами

        doc.save(docx_path)
        return True, None

    except Exception as e:
        return False, str(e)


def append_ep_stamp_to_pdf(pdf_path, signatures):
    """
    Добавляет страницу ЭП в конец существующего PDF.

    Args:
        pdf_path: абсолютный путь к .pdf файлу (будет перезаписан)
        signatures: QuerySet или список ElectronicSignature
    """
    if not os.path.exists(pdf_path):
        return False, 'Файл документа не найден'
    rows = _build_sig_rows(signatures)
    if not rows:
        return True, None

    try:
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from pypdf import PdfReader, PdfWriter

        W, H = A4
        buf = BytesIO()
        c = Canvas(buf, pagesize=A4)

        GREEN  = (0.024, 0.306, 0.243)
        GREY   = (0.42,  0.44,  0.50)
        BLACK  = (0.07,  0.08,  0.09)
        RED    = (0.60,  0.11,  0.11)
        LGREEN = (0.94,  0.99,  0.97)

        for idx, r in enumerate(rows):
            if idx > 0:
                c.showPage()

            # фон
            c.setFillColorRGB(*LGREEN)
            c.rect(1.5*cm, 2*cm, W - 3*cm, H - 4*cm, fill=1, stroke=0)

            # рамка
            c.setStrokeColorRGB(*GREEN)
            c.setLineWidth(1.5)
            c.rect(1.5*cm, 2*cm, W - 3*cm, H - 4*cm, fill=0, stroke=1)

            # заголовок
            c.setFillColorRGB(*GREEN)
            c.setFont('Helvetica-Bold', 11)
            c.drawCentredString(W / 2, H - 3*cm, 'КВАЛИФИЦИРОВАННАЯ ЭЛЕКТРОННАЯ ПОДПИСЬ')

            # статус
            is_valid = r['is_valid']
            c.setFont('Helvetica-Bold', 9)
            c.setFillColorRGB(*(GREEN if is_valid else RED))
            status = '✔  Подпись действительна' if is_valid else '✗  Подпись недействительна'
            c.drawCentredString(W / 2, H - 3.8*cm, status)

            # SVG-like cursive line (ломаная)
            c.setStrokeColorRGB(*GREEN)
            c.setLineWidth(1.8)
            pts = [
                (3.5*cm, H-5.2*cm), (5*cm, H-4.4*cm), (6*cm, H-5.5*cm),
                (7.2*cm, H-4.6*cm), (8.5*cm, H-5.3*cm), (10*cm, H-4.5*cm),
                (11.5*cm, H-5.4*cm), (13*cm, H-4.7*cm), (14.5*cm, H-5.2*cm),
                (16*cm, H-4.9*cm), (17.5*cm, H-5.3*cm),
            ]
            p = c.beginPath()
            p.moveTo(*pts[0])
            for px, py in pts[1:]:
                p.lineTo(px, py)
            c.drawPath(p, stroke=1, fill=0)
            # подчёркивание
            c.setLineWidth(0.5)
            c.setStrokeColorRGB(0.3, 0.7, 0.5)
            c.line(3.5*cm, H-5.5*cm, 17.5*cm, H-5.5*cm)

            # таблица полей
            data = [
                ('Подписант', r['full_name']),
            ]
            if r['position']:       data.append(('Должность', r['position']))
            if r['organization']:   data.append(('Организация', r['organization']))
            if r['inn']:            data.append(('ИНН', r['inn']))
            if r['snils']:          data.append(('СНИЛС', r['snils']))
            data.append(('Дата подписания', r['signed_at']))
            if r['ip']:             data.append(('IP адрес', r['ip']))
            data.append(('Алгоритм подписи', r['algorithm']))
            if r['issuer']:         data.append(('Издатель (УЦ)', r['issuer']))
            if r['serial']:         data.append(('Серийный номер', r['serial']))
            if r['valid_from'] and r['valid_to']:
                data.append(('Срок действия', f"{r['valid_from']} — {r['valid_to']}"))
            if r['thumbprint']:     data.append(('Отпечаток SHA-1', r['thumbprint']))
            data.append(('Хеш документа (SHA-256)', r['hash']))

            y = H - 6.3*cm
            lx, vx = 2.2*cm, 7.5*cm
            row_h = 0.65*cm

            for label, value in data:
                c.setFont('Helvetica', 7.5)
                c.setFillColorRGB(*GREY)
                c.drawString(lx, y, label + ':')
                c.setFont('Helvetica-Bold' if label == 'Подписант' else 'Helvetica', 8)
                c.setFillColorRGB(*BLACK)
                # длинные строки: укорачиваем
                max_w = W - vx - 1.5*cm
                while c.stringWidth(value, 'Helvetica-Bold' if label == 'Подписант' else 'Helvetica', 8) > max_w and len(value) > 10:
                    value = value[:-4] + '...'
                c.drawString(vx, y, value)
                y -= row_h
                if y < 2.5*cm:
                    break

            c.showPage()

        c.save()
        buf.seek(0)
        sig_pdf = PdfReader(buf)

        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        for page in sig_pdf.pages:
            writer.add_page(page)

        tmp_path = pdf_path + '.tmp'
        with open(tmp_path, 'wb') as f:
            writer.write(f)
        os.replace(tmp_path, pdf_path)
        return True, None

    except Exception as e:
        return False, str(e)

